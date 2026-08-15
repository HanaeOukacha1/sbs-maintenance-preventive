"""
Service d'export des rapports de maintenance.
Utilise les fichiers MASTER DATA originaux comme templates.
Principe : charger le template exact, injecter les données dynamiques, retourner le fichier.
"""
import os
import re
import tempfile
from io import BytesIO
from copy import copy
from datetime import date

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx2pdf import convert
import openpyxl
from openpyxl.drawing.image import Image as XlImage

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STATIC_DIR = os.path.join(BASE_DIR, "static")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")


# =============================================================================
# HELPERS
# =============================================================================

def convert_docx_to_pdf_buffer(docx_buffer: BytesIO) -> BytesIO:
    """Convertit un buffer DOCX en PDF via MS Word (nécessite Word installé)."""
    import pythoncom
    pythoncom.CoInitialize()  # Requis pour utiliser COM dans un thread (FastAPI)
    
    pdf_buffer = BytesIO()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_buffer.getvalue())
        tmp_docx = tmp.name
    tmp_pdf = tmp_docx.replace(".docx", ".pdf")
    try:
        convert(tmp_docx, tmp_pdf)
        with open(tmp_pdf, "rb") as f:
            pdf_buffer.write(f.read())
        pdf_buffer.seek(0)
    finally:
        for p in [tmp_docx, tmp_pdf]:
            if os.path.exists(p):
                try: os.remove(p)
                except: pass
        pythoncom.CoUninitialize()
    return pdf_buffer

def enforce_a4_landscape_docx(doc):
    """Force le format A4 Paysage (Landscape) pour donner plus d'espace en largeur."""
    from docx.shared import Mm
    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        # Réduire les marges pour maximiser l'espace
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

def enforce_a4_docx(doc):
    """Force le format A4 sur toutes les sections d'un document Word en modifiant l'XML."""
    from docx.oxml.ns import qn
    from docx.shared import Mm
    for section in doc.sections:
        if section.page_width > section.page_height:
            section.page_width = Mm(297)
            section.page_height = Mm(210)
            orient = 'landscape'
        else:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            orient = 'portrait'
        
        sectPr = section._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            if orient == 'landscape':
                pgSz.set(qn('w:w'), '16838')
                pgSz.set(qn('w:h'), '11906')
                pgSz.set(qn('w:orient'), 'landscape')
            else:
                pgSz.set(qn('w:w'), '11906')
                pgSz.set(qn('w:h'), '16838')
                pgSz.set(qn('w:orient'), 'portrait')


def get_logo_path(marche):
    """Retourne le chemin absolu du logo du client depuis le marché."""
    if marche and marche.logo_url:
        rel = marche.logo_url.strip("/")
        path = os.path.join(STATIC_DIR, *rel.replace("static/", "").split("/"))
        if os.path.exists(path):
            return path
    return None


def get_substancium_logo_path():
    """Retourne le chemin absolu du logo Substancium."""
    path = os.path.join(STATIC_DIR, "substancium_logo.png")
    return path if os.path.exists(path) else None


def get_eq_value(eq, intervention, field):
    """Récupère la valeur d'un champ: priorité aux données saisies par le technicien (equipement_modifie)."""
    base = getattr(eq, field, None) or ""
    if not intervention:
        return base
    reponses = intervention.reponses or {}
    # Priorité 1 : equipement_modifie (données saisies via la fiche mobile)
    modified = reponses.get("equipement_modifie", {})
    if isinstance(modified, dict) and modified.get(field):
        return modified[field]
    # Priorité 2 : champs hors inventaire
    if intervention.est_hors_inventaire and intervention.equipement_hors_inventaire:
        hi = intervention.equipement_hors_inventaire
        if isinstance(hi, dict) and hi.get(field):
            return hi[field]
    return base


def get_etat(intervention):
    """Retourne l'état/statut issu des réponses checklist du technicien."""
    if not intervention:
        return ""
    reponses = intervention.reponses or {}
    # Chercher dans l'ordre de priorité les clés possibles
    for key in ["statut", "etat", "observation", "etat_general", "etat_hardware", "etat_msante"]:
        val = reponses.get(key)
        if val:
            return str(val)
    return ""


def get_verifie(intervention):
    """Retourne OUI si le technicien a diagnostiqué cet équipement, NON sinon."""
    if not intervention or not (intervention.reponses):
        return "NON"
    reponses = intervention.reponses or {}
    # Considéré diagnostiqué si au moins une clé utile est remplie
    useful_keys = {"statut", "etat", "observation", "etat_general", "etat_hardware", "equipement_modifie"}
    if any(reponses.get(k) for k in useful_keys):
        return "OUI"
    return "NON"


def get_observation(intervention):
    """Retourne l'observation du technicien."""
    if not intervention:
        return ""
    reponses = intervention.reponses or {}
    return reponses.get("observation") or reponses.get("notes") or reponses.get("commentaire") or ""


def add_logos_to_excel(sheet, marche, logo_cell="A1", subs_cell="E1", client_w=100, client_h=50, subs_w=100, subs_h=50):
    """Injecte les deux logos dans une feuille Excel avec des tailles personnalisables."""
    client_logo = get_logo_path(marche)
    subs_logo = get_substancium_logo_path()
    if client_logo:
        try:
            img = XlImage(client_logo)
            img.width, img.height = client_w, client_h
            sheet.add_image(img, logo_cell)
        except Exception:
            pass
    if subs_logo:
        try:
            img2 = XlImage(subs_logo)
            img2.width, img2.height = subs_w, subs_h
            sheet.add_image(img2, subs_cell)
        except Exception:
            pass


def clean_workbook_for_export(wb):
    """
    Supprime les tableaux croisés dynamiques (pivots), validations de données
    et les Tableaux Excel structurés (ListObjects) pour éviter la corruption du fichier par openpyxl.
    """
    for ws in wb.worksheets:
        ws._pivots = []
        if hasattr(ws, 'data_validations'):
            ws.data_validations.dataValidation = []
        # Clear Excel Tables to prevent corruption when appending rows
        if hasattr(ws, '_tables'):
            from openpyxl.worksheet.table import TableList
            ws._tables = TableList()
        
        # Openpyxl doesn't fully remove pivot caches from workbook, but we do our best.


def copy_row_style(sheet, src_row, dest_row):
    """Copie le style d'une ligne source vers une ligne destination."""
    for col in range(1, sheet.max_column + 1):
        src_cell = sheet.cell(row=src_row, column=col)
        dest_cell = sheet.cell(row=dest_row, column=col)
        if src_cell.has_style:
            dest_cell.font = copy(src_cell.font)
            dest_cell.border = copy(src_cell.border)
            dest_cell.fill = copy(src_cell.fill)
            dest_cell.alignment = copy(src_cell.alignment)
            dest_cell.number_format = src_cell.number_format


def find_header_row(sheet, keywords=("N°", "Désignation", "Type", "DESIGNATION", "TYPE", "UTILISATEUR")):
    """Trouve la ligne d'en-tête des équipements (la 1ère ligne contenant un des mots-clés)."""
    for row_idx in range(1, min(sheet.max_row + 1, 20)):
        for col_idx in range(1, min(sheet.max_column + 1, 10)):
            val = sheet.cell(row=row_idx, column=col_idx).value
            if val:
                v_upper = str(val).upper()
                if "MARCHE" in v_upper or "MARCHÉ" in v_upper:
                    continue
                if any(kw in v_upper for kw in [k.upper() for k in keywords]):
                    return row_idx
    return None


# =============================================================================
# EXPORT ANCFCC (Word → PDF)
# Structure: 4 tableaux
#   Table 0: Client/Marché + Date
#   Table 1: Signatures (Technicien / Responsable)
#   Table 2: Infos UPS (Zone, Ville, Puissance, Batteries, Marque/Modèle)
#   Table 3: Checklist 10 points (oui / observations)
# =============================================================================

def export_ancfcc(mission, interventions, equipements):
    template_path = os.path.join(TEMPLATES_DIR, "template_ancfcc.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template ANCFCC introuvable: {template_path}")

    from docx import Document
    doc_base = Document(template_path)

    marche = mission.site.marche if mission.site else None
    site = mission.site
    intervention = interventions[0] if interventions else None
    eq = next((e for e in equipements if intervention and e.id == intervention.equipement_id), 
              equipements[0] if equipements else None)

    reponses = (intervention.reponses or {}) if intervention else {}
    numero_marche = (marche.numero if marche else None) or "N°021/2024/ANCFCC/DLA"
    site_nom = site.nom if site else ""
    ville = site.ville if site else ""
    date_str = intervention.date_intervention.strftime("%d / %m / %Y") if (intervention and intervention.date_intervention) else "..../.…/……"
    nom_technicien = (intervention.signature_technicien or "") if intervention else ""
    nom_responsable = (intervention.signature_client or "") if intervention else ""

    # Table 0 : Client/Marché + Date
    t0 = doc_base.tables[0]
    t0.cell(0, 1).text = f"ANCFCC / {numero_marche}"
    t0.cell(1, 1).text = date_str

    # Table 1 : Signatures
    t1 = doc_base.tables[1]
    t1.cell(1, 0).text = nom_technicien
    t1.cell(1, 1).text = nom_responsable

    # Table 2 : Infos UPS
    if len(doc_base.tables) > 2:
        t2 = doc_base.tables[2]
        # Ville
        t2.cell(2, 1).text = ville
        # Puissance KVA
        puissance = get_eq_value(eq, intervention, "puissance_kva") if eq else reponses.get("puissance_kva", "")
        t2.cell(0, 3).text = f"{puissance} KVA" if puissance else "…….. KVA"
        # Batteries
        nb_batt = get_eq_value(eq, intervention, "nb_batteries") if eq else reponses.get("nb_batteries", "")
        if len(t2.rows) > 1 and len(t2.columns) > 3:
            t2.cell(1, 3).text = str(nb_batt) if nb_batt else ""
        # Marque/Modèle
        if eq:
            marque_modele = f"{get_eq_value(eq, intervention, 'marque')} {get_eq_value(eq, intervention, 'modele')}".strip()
            if len(t2.rows) > 2 and len(t2.columns) > 3:
                t2.cell(2, 3).text = marque_modele or t2.cell(2, 3).text

    # Table 3 : Checklist 10 points
    if len(doc_base.tables) > 3:
        t3 = doc_base.tables[3]
        reponses = (interventions[0].reponses or {}) if interventions else {}
        for i in range(10):
            # Les clés peuvent être des strings "0".."9" ou int 0..9 ou "pt1".."pt10"
            rep = (reponses.get(str(i)) or
                   reponses.get(i) or
                   reponses.get(f"pt{i+1}") or
                   reponses.get(f"point_{i+1}") or {})
            if isinstance(rep, dict):
                reponse_val = str(rep.get("reponse", "")).lower()
                obs = rep.get("observation") or ""
            else:
                reponse_val = str(rep).lower()
                obs = ""
            oui = "X" if reponse_val in ["oui", "ok", "fait", "bon", "x"] else ""
            row_idx = i + 1
            if row_idx < len(t3.rows):
                if len(t3.columns) > 1:
                    t3.cell(row_idx, 1).text = oui
                if len(t3.columns) > 2:
                    t3.cell(row_idx, 2).text = obs

    # Enforce A4 format
    enforce_a4_docx(doc_base)
    
    # Sauvegarder en buffer
    buffer = BytesIO()
    doc_base.save(buffer)
    buffer.seek(0)
    pdf_buffer = convert_docx_to_pdf_buffer(buffer)
    return pdf_buffer, "application/pdf", f"Rapport_ANCFCC_{site_nom}.pdf"


# =============================================================================
# EXPORT ADM (Word → PDF)
# Structure: paragraphes d'en-tête + 1 tableau d'équipements (serveurs)
# Colonnes: Désignation | Fabricant | Modèle | N° Série | Processeur | Mémoire | DD-C | DD-D | IP | SW | HW
# =============================================================================

def export_adm(mission, interventions, equipements):
    template_path = os.path.join(TEMPLATES_DIR, "template_adm.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template ADM introuvable: {template_path}")

    from docx import Document
    doc_base = Document(template_path)

    marche = mission.site.marche if mission.site else None
    site = mission.site
    numero_marche = (marche.numero if marche else None) or "M0103/25"
    site_nom = site.nom if site else ""
    date_debut = ""
    date_fin = ""
    nom_technicien = ""
    if interventions:
        first = interventions[0]
        if first.date_intervention:
            date_debut = first.date_intervention.strftime("%d/%m/%Y")
        nom_technicien = first.signature_technicien or ""

    # Modifier les paragraphes d'en-tête
    for para in doc_base.paragraphs:
        if "Date et heure début" in para.text:
            for run in para.runs:
                if "….." in run.text or "……" in run.text:
                    run.text = run.text.replace("..../.…/……        …..h…..", date_debut + "        ")
        if "Marché N°" in para.text or "March" in para.text:
            for run in para.runs:
                if "M0103" in run.text or "March" in run.text:
                    run.text = re.sub(r"March[ée]\s*N°[\w/\.]+", f"Marché N°{numero_marche}", run.text)
        if "Site" in para.text and site_nom and (":…" in para.text or ": …" in para.text):
            for run in para.runs:
                run.text = re.sub(r"Site\s*:[\s…\.]+", f"Site : {site_nom}", run.text)
        if "Nom du Technicien" in para.text and nom_technicien:
            for run in para.runs:
                if "……" in run.text:
                    run.text = run.text.replace("…………………………………………..","  " + nom_technicien)

    # Remplir le tableau d'équipements
    if doc_base.tables:
        table = doc_base.tables[0]
        header_row = 0  # La ligne 0 est l'en-tête

        # Effacer les lignes de données existantes (garder uniquement l'en-tête)
        # En python-docx, on met à jour cellule par cellule
        data_start = 1
        equip_list = []
        for eq in equipements:
            intervention = next((i for i in interventions if i.equipement_id == eq.id), None)
            reponses = (intervention.reponses or {}) if intervention else {}
            eq_mod = reponses.get("equipement_modifie", {}) if reponses else {}

            def get_f(field):
                return eq_mod.get(field) or getattr(eq, field, "") or ""

            etat_sw = reponses.get("etat_software") or reponses.get("etat_sw") or reponses.get("etat") or reponses.get("statut") or ""
            etat_hw = reponses.get("etat_hardware") or reponses.get("etat_hw") or ""
            obs = get_observation(intervention)
            equip_list.append({
                "designation": get_f("designation") or getattr(eq, "type_equipement", "") or getattr(eq, "famille", ""),
                "fabricant":   get_f("marque"),
                "modele":      get_f("modele"),
                "n_serie":     get_f("numero_serie"),
                "processeur":  get_f("cpu") or getattr(eq, "processeur", "") or "",
                "memoire":     get_f("ram") or getattr(eq, "memoire", "") or "",
                "dd_c":        get_f("disque_c") or getattr(eq, "disque_c", "") or get_f("disque_dur") or "",
                "dd_d":        get_f("disque_d") or getattr(eq, "disque_d", "") or "",
                "ip":          get_f("ip") or getattr(eq, "adresse_ip", "") or getattr(eq, "ip", "") or "",
                "etat_sw":     etat_sw,
                "etat_hw":     etat_hw,
                "verifie":     get_verifie(intervention),
            })

        # Mettre à jour les lignes existantes du tableau
        for i, eq_data in enumerate(equip_list):
            row_idx = data_start + i
            if row_idx >= len(table.rows):
                break  # Ne pas dépasser les lignes du template
            row = table.rows[row_idx]
            vals = [
                eq_data["designation"], eq_data["fabricant"], eq_data["modele"],
                eq_data["n_serie"], eq_data["processeur"], eq_data["memoire"],
                eq_data["dd_c"], eq_data["dd_d"], eq_data["ip"],
                eq_data["etat_sw"], eq_data["etat_hw"]
            ]
            for ci, val in enumerate(vals):
                if ci < len(row.cells):
                    row.cells[ci].text = str(val) if val else row.cells[ci].text

    # Forcer le format Paysage pour avoir plus de place en largeur
    enforce_a4_landscape_docx(doc_base)
    
    buffer = BytesIO()
    doc_base.save(buffer)
    buffer.seek(0)
    pdf_buffer = convert_docx_to_pdf_buffer(buffer)
    return pdf_buffer, "application/pdf", f"Rapport_ADM_{site_nom}.pdf"


# =============================================================================
# EXPORT ANP (Excel multi-feuilles)
# Structure par feuille: 
#   Rows 1-3: En-tête (Agence, Master Data, Marché)
#   Row 4: vide
#   Row 5: Période
#   Row 6: Site
#   Row 7: En-têtes colonnes (N°, Désignation, Marque, Modèle, N° série, Etat)
#   Rows 8+: Données équipements
# =============================================================================

def export_anp(mission, interventions, equipements):
    template_path = os.path.join(TEMPLATES_DIR, "template_anp.xlsx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template ANP introuvable: {template_path}")

    wb = openpyxl.load_workbook(template_path)
    clean_workbook_for_export(wb)
    marche = mission.site.marche if mission.site else None
    site = mission.site
    site_nom = site.nom if site else ""
    numero_marche = (marche.numero if marche else None) or "01/DPJR-ANP/2025"
    today = date.today()
    periode = f"{today.strftime('%d/%m/%Y')}"

    # Identifier la feuille correspondant au site
    target_sheet = None
    for sname in wb.sheetnames:
        if site_nom and site_nom.upper() in sname.upper():
            target_sheet = wb[sname]
            break
    if target_sheet is None:
        target_sheet = wb.active

    # Mettre à jour l'en-tête
    target_sheet.cell(row=3, column=1).value = f"MARCHE CADRE N° {numero_marche}"
    target_sheet.cell(row=5, column=1).value = "Période : "   # Laissé vide — à remplir après impression
    target_sheet.cell(row=6, column=1).value = f"Site: {site_nom}"

    # Logos
    add_logos_to_excel(target_sheet, marche, logo_cell="G1", subs_cell="H1", w=80, h=40)

    # Trouver la ligne d'en-tête et effacer les données existantes
    header_row = find_header_row(target_sheet) or 7
    data_start = header_row + 1

    # Effacer toutes les lignes de données existantes
    for row_idx in range(data_start, target_sheet.max_row + 1):
        for col_idx in range(1, 7):
            target_sheet.cell(row=row_idx, column=col_idx).value = None

    # Remplir avec les données des interventions
    interv_map = {i.equipement_id: i for i in interventions}
    for idx, eq in enumerate(equipements):
        row_idx = data_start + idx
        intervention = interv_map.get(eq.id)

        designation = get_eq_value(eq, intervention, "designation") or getattr(eq, "type_equipement", "") or getattr(eq, "famille", "")
        marque = get_eq_value(eq, intervention, "marque")
        modele = get_eq_value(eq, intervention, "modele")
        n_serie = get_eq_value(eq, intervention, "numero_serie")
        etat = get_etat(intervention)
        verifie = get_verifie(intervention)

        copy_row_style(target_sheet, data_start, row_idx)

        target_sheet.cell(row=row_idx, column=1).value = idx + 1
        target_sheet.cell(row=row_idx, column=2).value = designation
        target_sheet.cell(row=row_idx, column=3).value = marque
        target_sheet.cell(row=row_idx, column=4).value = modele
        target_sheet.cell(row=row_idx, column=5).value = n_serie
        target_sheet.cell(row=row_idx, column=6).value = etat
        target_sheet.cell(row=row_idx, column=7).value = verifie

    # Zone de signature à la fin
    sig_row = data_start + len(equipements) + 2
    try:
        target_sheet.cell(row=sig_row, column=1).value = "Signature Technicien :"
        target_sheet.cell(row=sig_row + 4, column=1).value = "Signature Client / Responsable :"
    except Exception:
        pass

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"Rapport_ANP_{site_nom}.xlsx"


# =============================================================================
# EXPORT AMEE (Excel multi-feuilles par type d'équipement)
# Chaque feuille = 1 type (DATA CENTER, UC, IMPRIMANTE ET MFP, etc.)
# En-tête: marche en colonne B/E, ligne 1-3. Données débutent après l'en-tête colonnes.
# =============================================================================

def export_amee(mission, interventions, equipements):
    marche = mission.site.marche if mission.site else None
    site = mission.site
    site_nom = site.nom if site else ""

    # Choisir le bon template selon le site (Marrakech vs Rabat)
    if site_nom and "RABAT" in site_nom.upper():
        template_path = os.path.join(TEMPLATES_DIR, "template_amee_rabat.xlsx")
    else:
        template_path = os.path.join(TEMPLATES_DIR, "template_amee_marrakech.xlsx")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template AMEE introuvable: {template_path}")

    wb = openpyxl.load_workbook(template_path)
    clean_workbook_for_export(wb)
    numero_marche = (marche.numero if marche else None) or "29/2025"

    # ─── Construire un dict intervention par equipement_id (dernière version) ───
    # Clé: equipement_id → Intervention
    interv_map = {}
    for i in interventions:
        if i.equipement_id:
            interv_map[i.equipement_id] = i

    # ─── Construire un dict equipement par id ────────────────────────────────
    eq_map = {eq.id: eq for eq in equipements}

    # ─── Regrouper les équipements par feuille AMEE ───────────────────────────
    # Priorité: utiliser la feuille de l'intervention si disponible,
    # sinon utiliser sous_site de l'équipement
    FEUILLE_SHEET_MAP = {
        "DATA CENTER":        "DATA CENTER",
        "UC":                 "UC ",
        "MISE A JOUR":        "MISE A JOUR",
        "AVANCEE":            "AVANCEE",
        "IMPRIMANTE ET MFP":  "IMPRIMANTE ET MFP ",
        "IMP ET MFP":         "IMPRIMANTE ET MFP ",
        "IMPRIMANTE":         "IMPRIMANTE ET MFP ",
    }

    # Regrouper tous les équipements par leur feuille
    equip_by_feuille = {}
    for eq in equipements:
        # Feuille déterminée par intervention.feuille ou sous_site
        interv = interv_map.get(eq.id)
        if interv and interv.feuille:
            feuille = interv.feuille.upper().strip()
        else:
            feuille = (getattr(eq, "sous_site", "") or "").upper().strip()
            if not feuille:
                feuille = "UC"
        equip_by_feuille.setdefault(feuille, []).append(eq)

    # ─── Remplir chaque feuille du classeur ──────────────────────────────────
    used_sheet_names = set()
    for feuille_key, eqs in equip_by_feuille.items():
        # Trouver le nom de la feuille Excel correspondante
        sheet_name = None
        for key, sname in FEUILLE_SHEET_MAP.items():
            if key in feuille_key or feuille_key in key:
                sheet_name = sname
                break
        if not sheet_name:
            sheet_name = wb.sheetnames[0]  # fallback

        if sheet_name not in wb.sheetnames:
            continue

        sheet = wb[sheet_name]

        # Mettre à jour entête (site, marché, période)
        for row_idx in range(1, 10):
            for col_idx in range(1, sheet.max_column + 1):
                val = sheet.cell(row=row_idx, column=col_idx).value
                if not val:
                    continue
                val_str = str(val)
                if "Période" in val_str or "PERIODE" in val_str.upper():
                    sheet.cell(row=row_idx, column=col_idx).value = "Période du : ___________"
                elif "SITE" in val_str.upper() and ":" in val_str:
                    sheet.cell(row=row_idx, column=col_idx).value = f"SITE : {site_nom}"
                elif "MARCHE" in val_str.upper() and ("N°" in val_str or "N°" in val_str):
                    sheet.cell(row=row_idx, column=col_idx).value = f"MARCHÉ N° {numero_marche}"

        # Ajouter logos
        add_logos_to_excel(sheet, marche, logo_cell="A1", subs_cell="I1", client_w=80, client_h=40, subs_w=80, subs_h=40)

        # Trouver la ligne d'en-tête des colonnes
        header_row = find_header_row(sheet)
        if not header_row:
            continue
        data_start = header_row + 1

        # Lire les en-têtes de colonnes
        headers = {}
        for col_idx in range(1, sheet.max_column + 1):
            val = sheet.cell(row=header_row, column=col_idx).value
            if val:
                h = str(val).upper().strip()
                headers[h] = col_idx

        # Effacer les anciennes données
        for r in range(data_start, sheet.max_row + 1):
            for c in range(1, sheet.max_column + 1):
                sheet.cell(row=r, column=c).value = None

        # Conserver le nom des feuilles modifiées pour le nettoyage de fin
        used_sheet_names.add(sheet_name)

        # Helper: écrire dans une colonne par mot-clé
        def set_col(row_i, header_key, value):
            for h, c in headers.items():
                if header_key in h:
                    sheet.cell(row=row_i, column=c).value = value
                    return

        # ── Remplir toutes les lignes (diagnostiqués ET non diagnostiqués) ──
        for idx, eq in enumerate(eqs):
            row_i = data_start + idx
            copy_row_style(sheet, data_start - 1, row_i)

            interv = interv_map.get(eq.id)
            reponses = interv.reponses if interv else {}
            eq_mod = reponses.get("equipement_modifie", {}) if reponses else {}

            # Champs de l'équipement — priorité aux données saisies par le technicien
            def get_field(field):
                return eq_mod.get(field) or getattr(eq, field, "") or ""

            designation = get_field("designation") or getattr(eq, "famille", "") or getattr(eq, "type_equipement", "")
            marque_val  = get_field("marque")
            modele_val  = get_field("modele")
            sn_val      = get_field("numero_serie")
            utilisateur = get_field("utilisateur_nom")
            cpu_val     = get_field("cpu")
            ram_val     = get_field("ram")
            os_val      = get_field("systeme_exploitation")
            inv_val     = get_field("numero_inventaire")
            emp_val     = get_field("emplacement") or get_field("bureau")

            # Combiner Marque et Modèle si la colonne MARQUE n'existe pas dans le template (ex: Imprimantes)
            if marque_val and not any("MARQUE" in h for h in headers.keys()):
                modele_val = f"{marque_val} {modele_val}".strip()

            # Statut/État issu de la checklist
            if reponses:
                statut_val = (
                    reponses.get("statut") or
                    reponses.get("etat") or
                    reponses.get("observation") or
                    reponses.get("etat_general") or ""
                )
                verifie = "OUI"
            else:
                statut_val = ""
                verifie = "NON"

            # Numéro de ligne
            set_col(row_i, "N°", idx + 1)
            set_col(row_i, "NUM", idx + 1)

            # Champs communs
            set_col(row_i, "UTILISATEUR", utilisateur)
            set_col(row_i, "PERSONNE", utilisateur)
            set_col(row_i, "NOM", utilisateur)
            set_col(row_i, "TYPE", designation)
            set_col(row_i, "DÉSIGNATION", designation)
            set_col(row_i, "DESIGNATION", designation)
            set_col(row_i, "FAMILLE", designation)
            set_col(row_i, "MARQUE", marque_val)
            set_col(row_i, "MODÈLE", modele_val)
            set_col(row_i, "MODELE", modele_val)
            set_col(row_i, "S/N", sn_val)
            set_col(row_i, "SN", sn_val)
            set_col(row_i, "SÉRIE", sn_val)
            set_col(row_i, "SERIE", sn_val)
            set_col(row_i, "ETAT", statut_val)
            set_col(row_i, "ÉTAT", statut_val)
            set_col(row_i, "STATUT", statut_val)
            set_col(row_i, "VÉRIFIÉ", verifie)
            set_col(row_i, "VERIFIE", verifie)
            set_col(row_i, "CPU", cpu_val)
            set_col(row_i, "RAM", ram_val)
            set_col(row_i, "OS", os_val)
            set_col(row_i, "SYSTÈME", os_val)
            set_col(row_i, "SYSTEME", os_val)
            set_col(row_i, "INVENTAIRE", inv_val)
            set_col(row_i, "INV", inv_val)
            set_col(row_i, "EMPLACEMENT", emp_val)
            set_col(row_i, "BUREAU", emp_val)
            set_col(row_i, "OBSERVATIONS", reponses.get("notes", "") if reponses else "")
            set_col(row_i, "NOTES", reponses.get("notes", "") if reponses else "")

        # ── Ajouter zone de signature à la fin ──────────────────────────────
        sig_row = data_start + len(eqs) + 2
        try:
            sheet.cell(row=sig_row, column=1).value = "Signature Technicien :"
            sheet.cell(row=sig_row + 4, column=1).value = "Signature Client / Responsable :"
        except Exception:
            pass

    # Nettoyage final : Supprimer les feuilles inutilisées (Feuil9, Feuil1, DATA CENTER si vide, etc.)
    for s_name in list(wb.sheetnames):
        if s_name not in used_sheet_names:
            try:
                wb.remove(wb[s_name])
            except Exception:
                pass

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"Rapport_AMEE_{site_nom.replace(' ', '_')}.xlsx"



# =============================================================================
# EXPORT GÉNÉRIQUE EXCEL (ONP, MARSA, AOH, MHAI, CNDH, INPPLC, MSANTE)
# Tous ont une structure similaire:
#   Lignes 1-N : en-tête (ne pas modifier sauf Période/Site/Marché)
#   1 ligne d'en-tête colonnes
#   Lignes suivantes : données équipements
# =============================================================================

def export_generic_excel(mission, interventions, equipements, template_filename, client_name):
    template_path = os.path.join(TEMPLATES_DIR, template_filename)
    if not os.path.exists(template_path):
        # Fallback: template simplifié
        return export_generic_fallback(mission, interventions, equipements, client_name)

    wb = openpyxl.load_workbook(template_path)
    clean_workbook_for_export(wb)
    marche = mission.site.marche if mission.site else None
    site = mission.site
    site_nom = site.nom if site else ""
    numero_marche = (marche.numero if marche else None) or ""
    # Période laissée vide — à remplir manuellement après impression

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        if sheet.max_row < 2:
            continue

        # Mettre à jour les cellules d'en-tête dynamiquement
        for row_idx in range(1, min(sheet.max_row + 1, 15)):
            for col_idx in range(1, min(sheet.max_column + 1, 5)):
                val = sheet.cell(row=row_idx, column=col_idx).value
                if not val:
                    continue
                val_str = str(val)
                if "Période" in val_str or "PERIODE" in val_str.upper():
                    sheet.cell(row=row_idx, column=col_idx).value = "Période : ___________"  # À remplir manuellement
                elif "Site:" in val_str or "SITE:" in val_str.upper():
                    sheet.cell(row=row_idx, column=col_idx).value = f"Site: {site_nom}"
                elif "MARCHE CADRE" in val_str.upper() and numero_marche:
                    sheet.cell(row=row_idx, column=col_idx).value = f"MARCHE CADRE N° {numero_marche}"
                elif "Affectation" in val_str:
                    sheet.cell(row=row_idx, column=col_idx).value = f"Affectation : {site_nom}"

        # Ajouter logos
        add_logos_to_excel(sheet, marche)

        # Trouver la ligne d'en-tête des colonnes
        header_row = find_header_row(sheet)
        if not header_row:
            continue
        data_start = header_row + 1

        # Lire les en-têtes des colonnes
        headers = {}
        for col_idx in range(1, sheet.max_column + 1):
            val = sheet.cell(row=header_row, column=col_idx).value
            if val:
                headers[str(val).upper().strip()] = col_idx

        # Effacer les données existantes
        ref_row = data_start
        for r in range(data_start, sheet.max_row + 1):
            for c in range(1, sheet.max_column + 1):
                sheet.cell(row=r, column=c).value = None

        # Écrire les données
        interv_map_g = {i.equipement_id: i for i in interventions}
        for idx, eq in enumerate(equipements):
            row_idx = data_start + idx
            intervention = interv_map_g.get(eq.id)
            copy_row_style(sheet, ref_row, row_idx)

            eq_mod = (intervention.reponses or {}).get("equipement_modifie", {}) if intervention and intervention.reponses else {}
            def get_f(field):
                return eq_mod.get(field) or getattr(eq, field, "") or ""

            designation = get_f("designation") or getattr(eq, "type_equipement", "") or getattr(eq, "famille", "")
            marque_val  = get_f("marque")
            modele_val  = get_f("modele")
            sn_val      = get_f("numero_serie")
            utilisateur = get_f("utilisateur_nom")
            etat_val    = get_etat(intervention)
            verifie_val = get_verifie(intervention)
            obs_val     = get_observation(intervention)

            def set_col(keys, value):
                for key in keys:
                    for h, c in headers.items():
                        if key in h:
                            sheet.cell(row=row_idx, column=c).value = value
                            return

            # Numéro ordinal
            if "N°" in headers or "NO" in headers:
                c = headers.get("N°") or headers.get("NO") or 1
                sheet.cell(row=row_idx, column=c).value = idx + 1
            else:
                sheet.cell(row=row_idx, column=1).value = idx + 1

            set_col(["DESIGNATION", "DÉSIGNATION", "TYPE", "DESCRIPTION", "FAMILLE"], designation)
            set_col(["MARQUE", "FABRICANT"], marque_val)
            set_col(["MODELE", "MODÈLE"], modele_val)
            set_col(["SERIE", "SN", "S/N", "N° SERIE"], sn_val)
            set_col(["ETAT", "ÉTAT", "STATUT"], etat_val)
            set_col(["OBSERVATION", "OBSERVATIONS", "REMARQUE", "NOTES"], obs_val)
            set_col(["UTILISATEUR", "PERSONNE", "NOM"], utilisateur)
            set_col(["VÉRIFIÉ", "VERIFIE"], verifie_val)

        # Zone de signature à la fin
        sig_row = data_start + len(equipements) + 2
        try:
            sheet.cell(row=sig_row, column=1).value = "Signature Technicien :"
            sheet.cell(row=sig_row + 4, column=1).value = "Signature Client / Responsable :"
        except Exception:
            pass

        break  # Traiter uniquement la 1ère feuille active pour les clients simples

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"Rapport_{client_name}_{site_nom}.xlsx"


def export_generic_fallback(mission, interventions, equipements, client_name):
    """Template minimal si le fichier de base n'existe pas."""
    wb = openpyxl.Workbook()
    sheet = wb.active
    marche = mission.site.marche if mission.site else None
    site_nom = mission.site.nom if mission.site else ""
    numero = (marche.numero if marche else "") or ""
    sheet["A1"] = f"{client_name} - MARCHE CADRE N° {numero}"
    sheet["A2"] = f"Site: {site_nom}"
    sheet["A3"] = f"Période: {date.today().strftime('%d/%m/%Y')}"
    headers = ["N°", "Désignation", "Marque", "Modèle", "N° Série", "Etat", "Observation"]
    for ci, h in enumerate(headers, 1):
        sheet.cell(row=5, column=ci).value = h
    for idx, eq in enumerate(equipements):
        intervention = next((i for i in interventions if i.equipement_id == eq.id), None)
        row = 6 + idx
        sheet.cell(row=row, column=1).value = idx + 1
        sheet.cell(row=row, column=2).value = getattr(eq, "type_equipement", "") or getattr(eq, "famille", "")
        sheet.cell(row=row, column=3).value = get_eq_value(eq, intervention, "marque")
        sheet.cell(row=row, column=4).value = get_eq_value(eq, intervention, "modele")
        sheet.cell(row=row, column=5).value = get_eq_value(eq, intervention, "numero_serie")
        sheet.cell(row=row, column=6).value = get_etat(intervention)
        sheet.cell(row=row, column=7).value = get_observation(intervention)
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"Rapport_{client_name}_{site_nom}.xlsx"


def export_dynamic_excel(mission, interventions, equipements):
    """
    Génère un classeur Excel dynamiquement à partir de zéro.
    - Colonnes de base : issues de l'équipement (equipement_modifie en priorité)
    - Colonnes formulaire : issues du JSON schema de la feuille (labels exacts de la BD)
    - Valeurs saisies par le technicien affichées, non-saisies = vide
    """
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    import json

    marche = mission.site.marche if mission.site else None
    marche_nom = marche.nom if marche else ""
    marche_num = marche.numero if marche else ""
    client_name = marche.client if marche else "Client"
    site_nom = mission.site.nom if mission.site else "Site"

    # --- Charger tous les JSON schemas depuis la DB (une seule fois) ---
    from app.db.database import SessionLocal
    from sqlalchemy import text
    db = SessionLocal()
    all_schemas_raw = db.execute(text("SELECT id, nom, schema_data FROM json_schemas")).fetchall()
    schema_by_id = {}
    for sid, snom, sdata in all_schemas_raw:
        if sdata:
            if isinstance(sdata, str):
                try:
                    sdata = json.loads(sdata)
                except Exception:
                    sdata = []
            schema_by_id[sid] = sdata  # liste de {key, label, ...}
    db.close()

    def get_schema_fields(schema_id):
        """Retourne la liste des champs du schema (key, label) en excluant les champs meta."""
        meta_keys = {"statut", "etat", "observation", "notes", "commentaire",
                     "etat_general", "etat_hardware", "equipement_modifie", "etat_msante",
                     "observation_cndh"}
        fields = []
        for f in schema_by_id.get(schema_id, []):
            if isinstance(f, dict) and f.get("key") not in meta_keys:
                fields.append({"key": f["key"], "label": f.get("label", f["key"])})
        return fields

    # --- Grouper les équipements par feuille ---
    equip_by_feuille = {}
    interv_map = {}  # equipement_id -> intervention
    # Si plusieurs interventions pour le même équipement, garder la plus récente
    for i in interventions:
        existing = interv_map.get(i.equipement_id)
        if existing is None or (i.updated_at and existing.updated_at and i.updated_at > existing.updated_at):
            interv_map[i.equipement_id] = i

    for eq in equipements:
        interv = interv_map.get(eq.id)
        feuille = "Equipements"
        if interv and interv.feuille:
            feuille = interv.feuille.strip()
        elif getattr(eq, "sous_site", ""):
            feuille = str(getattr(eq, "sous_site")).strip()
        equip_by_feuille.setdefault(feuille, []).append((eq, interv))

    # --- Build schema lookup by name too ---
    schema_by_nom = {}  # nom.upper() -> schema_id
    for sid, snom, sdata in all_schemas_raw:
        if snom:
            schema_by_nom[snom.upper()] = sid

    # --- Trouver le schema_id utilisé dans chaque feuille ---
    def get_feuille_schema_id(feuille_name, items):
        """Trouve le schema_id: priorité interventions, fallback recherche par nom de feuille."""
        # MSANTE utilise des formulaires en dur, pas de JSON schemas
        if mission.site and mission.site.checklist_type and mission.site.checklist_type.startswith('MSANTE'):
            return None
            
        for eq, interv in items:
            if interv and interv.json_schema_id:
                return interv.json_schema_id
        # Fallback: search schema whose name matches the feuille name
        fname_upper = feuille_name.upper().replace(" ", "_")
        for snom_upper, sid in schema_by_nom.items():
            if fname_upper == snom_upper:
                return sid
            if fname_upper in snom_upper or snom_upper in fname_upper:
                # Éviter que 'ADMIN' match 'ADM'
                if snom_upper == 'ADM' and fname_upper != 'ADM':
                    continue
                return sid
        return None

    # --- Styles Excel ---
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    BLUE_DARK = "1F4E78"
    header_fill = PatternFill(start_color=BLUE_DARK, end_color=BLUE_DARK, fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for feuille_name, items in equip_by_feuille.items():
        safe_title = "".join(c for c in feuille_name if c not in r'[]:*?/\ ')
        if not safe_title:
            safe_title = "Equipements"
        ws = wb.create_sheet(title=safe_title[:31])

        # ---- EN-TÊTE : Logos + Titre ----
        add_logos_to_excel(
            ws, marche,
            logo_cell="H1", subs_cell="A1",
            client_w=120, client_h=60,
            subs_w=120, subs_h=60
        )
        ws.row_dimensions[1].height = 50
        ws.row_dimensions[2].height = 25
        ws.row_dimensions[3].height = 20
        ws.row_dimensions[4].height = 25

        ws.merge_cells('C1:F1')
        ws['C1'].value = "Marché N° %s - %s" % (marche_num, marche_nom) if marche_nom else "Marché N° %s" % marche_num
        ws['C1'].font = Font(size=12, bold=True)
        ws['C1'].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        ws.merge_cells('C2:F2')
        ws['C2'].value = "Rapport d'activité"
        ws['C2'].font = Font(size=14, bold=True, color="004080")
        ws['C2'].alignment = Alignment(horizontal="center", vertical="center")

        ws.merge_cells('B4:F4')
        ws['B4'].value = "Période : ........................................................"
        ws['B4'].font = Font(size=11, bold=True)
        ws['B4'].alignment = Alignment(horizontal="left", vertical="center")

        # ---- COLONNES ----
        is_maj = "MAJ" in feuille_name.upper() or "MISE A JOUR" in feuille_name.upper()
        schema_id = get_feuille_schema_id(feuille_name, items)
        form_fields = get_schema_fields(schema_id) if schema_id else []

        # If schema didn't have specific form fields, fall back to dynamic keys from actual responses
        if not form_fields:
            meta_keys = {"statut", "etat", "observation", "notes", "commentaire",
                         "etat_general", "etat_hardware", "equipement_modifie",
                         "etat_msante", "observation_cndh"}
            seen = set()
            for eq, interv in items:
                if interv and interv.reponses:
                    for k, v in interv.reponses.items():
                        if k not in meta_keys and k not in seen:
                            seen.add(k)
                            form_fields.append({"key": k, "label": k.replace("_", " ").capitalize()})

        if is_maj:
            # Format spécial : Personnes | Types | [champs formulaire]
            base_cols = [
                {"key": "_personnes", "label": "Personnes"},
                {"key": "_types",     "label": "TYPES"},
            ]
        else:
            # Déterminer les champs d'équipement selon le marché (comme dans l'app mobile)
            checklist_type = mission.site.checklist_type if mission.site else ""
            
            base_cols = [
                {"key": "_num",              "label": "N°"},
                {"key": "designation",        "label": "Désignation"},
                {"key": "marque",             "label": "Marque"},
                {"key": "modele",             "label": "Modèle"},
                {"key": "numero_serie",       "label": "N° Série"},
            ]
            
            if checklist_type == 'AMEE_MARRAKECH':
                if "UC" in feuille_name.upper() or "PC" in feuille_name.upper():
                    base_cols.extend([
                        {"key": "utilisateur_nom", "label": "Utilisateur"},
                        {"key": "cpu", "label": "CPU"},
                        {"key": "ram", "label": "RAM"},
                        {"key": "systeme_exploitation", "label": "Système d'exploitation"},
                        {"key": "numero_inventaire", "label": "N° Inventaire"},
                    ])
            elif checklist_type == 'AMEE_RABAT':
                if "UC" in feuille_name.upper() or "PC" in feuille_name.upper():
                    base_cols.extend([
                        {"key": "utilisateur_nom", "label": "Utilisateur"},
                        {"key": "cpu", "label": "CPU"},
                        {"key": "ram", "label": "RAM"},
                        {"key": "systeme_exploitation", "label": "Système d'exploitation"},
                    ])
            elif checklist_type == 'MARSA_MAROC':
                base_cols.extend([
                    {"key": "direction", "label": "Direction"},
                    {"key": "bureau", "label": "Bureau"},
                    {"key": "utilisateur_nom", "label": "Utilisateur"},
                ])
            elif checklist_type in ['MHAI', 'ANP', 'AOH']:
                base_cols.extend([
                    {"key": "numero_inventaire", "label": "N° Inventaire"},
                ])
            elif checklist_type in ['MSANTE_CAPM', 'MSANTE_SIGNATURE'] or (checklist_type == 'MSANTE_DPRF' and 'COMPTABILITE' in feuille_name.upper()):
                base_cols.extend([
                    {"key": "utilisateur_nom", "label": "Utilisateur"},
                ])
            elif checklist_type and checklist_type.startswith('CNDH_'):
                # CNDH a "entite" au début
                base_cols.insert(1, {"key": "entite", "label": "Entité / Article"})
                if checklist_type == 'CNDH_G2':
                    base_cols.extend([
                        {"key": "emplacement", "label": "Emplacement"},
                        {"key": "affectation", "label": "Affectation"},
                    ])

        all_cols = base_cols + form_fields

        if not is_maj:
            all_cols += [
                {"key": "_etat",  "label": "État"},
                {"key": "_obs",   "label": "Observations"},
            ]

        # Écrire les en-têtes sur la ligne 5
        HEADER_ROW = 5
        DATA_START  = 6

        for col_idx, col in enumerate(all_cols, start=1):
            cell = ws.cell(row=HEADER_ROW, column=col_idx)
            cell.value = col["label"]
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = max(len(col["label"]) + 4, 14)

        # ---- DONNÉES ----
        for row_num, (eq, interv) in enumerate(items, start=DATA_START):
            reponses    = (interv.reponses or {}) if interv else {}
            eq_mod      = reponses.get("equipement_modifie", {}) or {}

            def cell_val(key):
                """Récupère la valeur pour une colonne donnée."""
                # Colonnes calculées
                if key == "_num":
                    return row_num - DATA_START + 1
                if key == "_personnes":
                    return eq_mod.get("utilisateur_nom") or getattr(eq, "utilisateur_nom", "") or ""
                if key == "_types":
                    d = eq_mod.get("designation") or getattr(eq, "designation", "") or ""
                    if not d:
                        t = getattr(eq, "type_equipement", "") or ""
                        if hasattr(t, "value"): t = t.value
                        f = getattr(eq, "famille", "") or ""
                        if hasattr(f, "value"): f = f.value
                        d = t or f
                    return d
                if key == "_etat":
                    for k in ["statut", "etat", "etat_general", "etat_hardware", "etat_msante", "observation", "observation_cndh"]:
                        v = reponses.get(k)
                        if v: return str(v)
                    return ""
                if key == "_obs":
                    return reponses.get("notes") or reponses.get("commentaire") or ""
                # Champs formulaire (clé directe dans reponses)
                if key in reponses:
                    return str(reponses[key]) if reponses[key] is not None else ""
                # Champs équipement (equipement_modifie en priorité, puis eq)
                v = eq_mod.get(key)
                if v:
                    return str(v)
                raw = getattr(eq, key, None)
                if raw is None: return ""
                if hasattr(raw, "value"): return raw.value
                return str(raw) if raw else ""

            for col_idx, col in enumerate(all_cols, start=1):
                c = ws.cell(row=row_num, column=col_idx)
                c.value = cell_val(col["key"])
                c.border = border
                c.alignment = Alignment(vertical="center", wrap_text=True)

        # ---- SIGNATURES ----
        # Pour MSANTE_DPRF, seules les feuilles "COMPTABILITE" ont une signature à la fin.
        if not (checklist_type == 'MSANTE_DPRF' and 'COMPTABILITE' not in feuille_name.upper()):
            sig_row = DATA_START + len(items) + 2
            ws.cell(row=sig_row, column=1).value = "Signature SBS (Technicien) :"
            ws.cell(row=sig_row, column=1).font = Font(bold=True)
            last_col = max(len(all_cols) - 1, 1)
            ws.cell(row=sig_row, column=last_col).value = "Signature Client / Responsable :"
            ws.cell(row=sig_row, column=last_col).font = Font(bold=True)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "Rapport_%s_%s.xlsx" % (client_name, site_nom)



# =============================================================================
# POINT D'ENTRÉE PRINCIPAL
# =============================================================================

def exporter_mission(mission, interventions, equipements):
    """Point d'entrée. Dispatche vers la bonne fonction selon le marché."""
    marche_nom = ""
    if mission.site and mission.site.marche:
        marche_nom = mission.site.marche.nom.upper()
    # Conserver uniquement ANCFCC et ADM pour le format Word
    if "ANCFCC" in marche_nom:
        return export_ancfcc(mission, interventions, equipements)
    if "ADM" in marche_nom:
        return export_adm(mission, interventions, equipements)

    # TOUS les autres rapports passent par le nouveau moteur dynamique 100% propre !
    return export_dynamic_excel(mission, interventions, equipements)
