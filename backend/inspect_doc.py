from docx import Document

def inspect_doc():
    doc = Document("C:/Users/hanae/Desktop/Stage PFA 2026/backend/app/templates/template_ancfcc.docx")
    print("--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for r, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"  Row {r}: {row_data}")

if __name__ == "__main__":
    inspect_doc()
