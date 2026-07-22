"""
DRY-RUN du seed SBS — Montre ce qui sera inséré sans toucher à la BDD.
Commande : python preview_seed.py
"""
import sys, os, re
from collections import defaultdict

try:
    import openpyxl
    import xlrd
    from docx import Document
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "openpyxl", "xlrd", "python-docx", "-q"])
    import openpyxl, xlrd
    from docx import Document

MASTER_DATA = os.path.normpath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "MASTER DATA")
)

# ========== COPY DE SEED.PY (fonctions) ==========

def mapper_type(article: str) -> str:
    a = str(article).upper().strip()
    if any(x in a for x in ["PORTABLE", "LAPTOP", "NOTEBOOK", " PT ", "PC PORT"]):
        return "PORTABLE"
    if any(x in a for x in ["UC", "UNITE CENTRALE", "DESKTOP", "PC BURE", "ORDINATEUR DE BUREAU"]):
        return "PC"
    if a in ["PC"] or a.startswith("PC "):
        return "PC"
    if any(x in a for x in ["SERVEUR", "SERVER", "RACK"]):
        return "SERVEUR"
    if any(x in a for x in ["ONDULEUR", "UPS", "ASC"]):
        return "ONDULEUR"
    if any(x in a for x in ["BAIE", "SWITCH", "ROUTEUR", "FIREWALL", "BRASSAGE", "STOCKAGE"]):
        return "BAIE_BRASSAGE"
    if any(x in a for x in ["IMPRIMANTE", "IMP ", "MFP", "LASER", "COPIEUR", "MULTIF"]):
        return "IMPRIMANTE"
    if any(x in a for x in ["ECRAN", "ÉCRAN", "MONITOR", "MONITEUR"]):
        return "ECRAN"
    if any(x in a for x in ["SCANNER", "SCAN "]):
        return "SCANNER"
    return "AUTRE"

def lire_excel(filepath: str) -> dict:
    ext = filepath.lower().rsplit(".", 1)[-1]
    result = {}
    try:
        if ext == "xlsx":
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            for sn in wb.sheetnames:
                ws = wb[sn]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    r = [("" if c is None else str(c).strip()) for c in row]
                    if any(r): rows.append(r)
                result[sn] = rows
            wb.close()
        else:
            wb = xlrd.open_workbook(filepath)
            for sn in wb.sheet_names():
                ws = wb.sheet_by_name(sn)
                rows = []
                for i in range(ws.nrows):
                    r = [str(ws.cell_value(i, j)).strip() for j in range(ws.ncols)]
                    if any(r): rows.append(r)
                result[sn] = rows
    except Exception as e:
        print(f"      ⚠️  {e}")
    return result

def extraire_depuis_feuille(rows: list) -> dict:
    header_idx = -1
    col = {}
    for i, row in enumerate(rows):
        r_up = [str(c).upper().strip() for c in row]
        found = 0
        tmp = {}
        for j, cell in enumerate(r_up):
            if cell in ["N° DE SERIE", "N° SERIE", "SN", "NUMÉRO DE SÉRIE", "NUMERO DE SERIE", "N°SERIE"]:
                tmp["serie"] = j; found += 1
            elif cell in ["MARQUE", "BRAND", "FABRICANT"]:
                tmp["marque"] = j; found += 1
            elif cell in ["MODELE", "MODÈLE", "MODEL", "MODÉLE"]:
                tmp["modele"] = j; found += 1
            elif cell in ["ARTICLE", "TYPE", "DÉSIGNATION", "DESIGNATION", "EQUIPEMENT", "MATÉRIEL", "MATERIEL"]:
                tmp["type"] = j; found += 1
            elif cell in ["AFFECTATION", "UTILISATEUR", "USER"]:
                tmp["utilisateur"] = j
            elif cell in ["OBSERVATION", "OBSERVATIONS", "STATUT", "ETAT", "ÉTAT"]:
                tmp["observation"] = j
            elif cell in ["SITE", "SITES", "LOCALISATION", "VILLE"]:
                tmp["site"] = j
        if found >= 2:
            header_idx = i
            col = tmp
            break

    if header_idx < 0 or not col:
        return {}

    par_site = defaultdict(list)
    has_site_col = "site" in col

    for row in rows[header_idx + 1:]:
        if not any(str(c).strip() for c in row):
            continue
        def get(key):
            idx = col.get(key)
            return str(row[idx]).strip() if idx is not None and idx < len(row) else ""

        type_val = get("type")
        marque_val = get("marque")
        modele_val = get("modele")
        serie_val = get("serie")
        site_val = get("site") if has_site_col else ""

        if not type_val or type_val.upper() in ["", "TYPE", "ARTICLE", "DÉSIGNATION", "DESIGNATION", "TOTAL", "SOUS-TOTAL", "NAN", "0.0"]:
            continue
        if not any([marque_val, modele_val, serie_val]):
            continue
        if serie_val.endswith(".0"): serie_val = serie_val[:-2]
        if serie_val.upper() in ["NAN", "N/A", "NONE", "", "0.0", "0"]: serie_val = None

        par_site[site_val].append({
            "nom": f"{type_val} {marque_val} {modele_val}".strip()[:80],
            "type": mapper_type(type_val),
            "serie": serie_val,
        })
    return dict(par_site)

def site_depuis_fichier(filename: str) -> str:
    nom = os.path.splitext(filename)[0]
    for pfx in ["NV MP ", "MP ", "MD ", "Masters Data ", "Checklist ANCFCC 132 ", "Checklist ANCFCC 132"]:
        if nom.upper().startswith(pfx.upper()):
            nom = nom[len(pfx):]
    nom = re.sub(r'\s+(S1|S2|1T|2T|3T|4T)[-_]?\d{0,4}.*$', '', nom, flags=re.IGNORECASE)
    nom = re.sub(r'\s+\d{4}$', '', nom)
    for tag in [" ok", " OK", " MODIFIER", " RESEAUX", " MODIFIER OK", " PS "]:
        nom = nom.replace(tag, "")
    return nom.strip()

# ========== PREVIEW ==========

MARCHES = {
    "ADM":         "Administration",
    "AMEE":        "Agence Marocaine pour l'Efficacité Énergétique",
    "ANCFCC":      "Agence Nationale de la Conservation Foncière",
    "ANP":         "Agence Nationale des Ports",
    "AOH":         "Al Omrane Holding",
    "CNDH":        "Conseil National des Droits de l'Homme",
    "INPPLC":      "Instance Nationale de la Probité et Lutte contre la Corruption",
    "MARSA MAROC": "Marsa Maroc",
    "MHAI":        "Ministère des Habous et des Affaires Islamiques",
    "MSANTE":      "Ministère de la Santé",
    "ONP":         "Office National des Pêches",
}

SITES_ANCFCC = [
    "AGADIR", "BENGUERIR", "KHEMISSAT", "KENITRA", "MDIEQ",
    "MIDELT", "OUARZAZATE", "ROMMANI", "SALA EL JADIDA", "SETTAT",
    "SIDI SLIMANE", "TETOUAN", "SIDI BANNOUR"
]

grand_total_sites = 0
grand_total_equips = 0
rapport_complet = {}

print("\n" + "="*70)
print("PREVIEW DU SEED — SBS Maintenance Préventive")
print("="*70)

for code, client in MARCHES.items():
    print(f"\n{'─'*70}")
    print(f"📁 {code} — {client}")
    dossier = os.path.join(MASTER_DATA, code)
    par_site = defaultdict(list)

    if os.path.isdir(dossier):
        fichiers = sorted(os.listdir(dossier))
        for fichier in fichiers:
            path = os.path.join(dossier, fichier)
            ext = fichier.lower().rsplit(".", 1)[-1]

            if ext == "docx":
                print(f"  📝 {fichier} (Word — checklist)")
                if code == "ANCFCC":
                    try:
                        doc = Document(path)
                        ville = ""
                        for table in doc.tables:
                            for row in table.rows:
                                cells = [c.text.strip() for c in row.cells]
                                for i, c in enumerate(cells):
                                    if c.upper() == "VILLE" and i+1 < len(cells) and cells[i+1].strip():
                                        ville = cells[i+1].strip()
                        if not ville:
                            ville = site_depuis_fichier(fichier)
                        if ville:
                            par_site[ville].append({"nom": "Onduleur Riello", "type": "ONDULEUR", "serie": None})
                    except: pass
                elif code == "ADM":
                    try:
                        doc = Document(path)
                        for table in doc.tables:
                            rows_t = [[c.text.strip() for c in row.cells] for row in table.rows]
                            hdr = False
                            col_des = -1
                            for i, r in enumerate(rows_t):
                                r_up = [c.upper() for c in r]
                                if "DÉSIGNATION" in r_up or "DESIGNATION" in r_up:
                                    col_des = next((j for j, c in enumerate(r_up) if "DÉSIGNATION" in c or "DESIGNATION" in c), -1)
                                    hdr = True; continue
                                if hdr and col_des >= 0 and col_des < len(r) and r[col_des]:
                                    par_site["Siège ADM"].append({"nom": r[col_des], "type": "SERVEUR", "serie": None})
                    except: pass
                continue

            if ext not in ("xlsx", "xls"):
                continue

            print(f"  📊 {fichier}")
            feuilles = lire_excel(path)

            if code in ("CNDH", "MSANTE", "AMEE"):
                nom_site = site_depuis_fichier(fichier)
                for nom_feuille, rows in feuilles.items():
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        s = site_col if site_col else nom_site
                        par_site[s].extend(equips)

            elif code in ("ANP",):
                for nom_feuille, rows in feuilles.items():
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        s = site_col if site_col else nom_feuille
                        par_site[s].extend(equips)

            elif code in ("ONP", "MARSA MAROC", "AOH", "INPPLC", "MHAI"):
                for nom_feuille, rows in feuilles.items():
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        if site_col:
                            par_site[site_col].extend(equips)
                        elif nom_feuille.lower() not in ["feuil1", "feuil2", "feuil3", "sheet1"]:
                            par_site[nom_feuille.strip()].extend(equips)
                        else:
                            par_site[f"Siège {code}"].extend(equips)
            else:
                for nom_feuille, rows in feuilles.items():
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        s = site_col if site_col else nom_feuille
                        par_site[s].extend(equips)

    # ANCFCC : garantir les 13 sites
    if code == "ANCFCC":
        for v in SITES_ANCFCC:
            if v not in par_site:
                par_site[v] = []

    if not par_site:
        par_site[f"Siège {code}"] = []

    nb_sites = len(par_site)
    nb_equips = sum(len(v) for v in par_site.values())
    grand_total_sites += nb_sites
    grand_total_equips += nb_equips
    rapport_complet[code] = par_site

    print(f"\n  → {nb_sites} site(s) | {nb_equips} équipement(s)")
    for nom_site, equips in sorted(par_site.items()):
        print(f"    📍 {nom_site} ({len(equips)} équip.)")
        # Compter par type
        types_count = defaultdict(int)
        for e in equips:
            types_count[e["type"]] += 1
        if types_count:
            detail = " | ".join(f"{t}:{n}" for t, n in sorted(types_count.items()))
            print(f"       [{detail}]")

print(f"\n{'='*70}")
print("RÉSUMÉ GLOBAL")
print(f"{'='*70}")
print(f"  Marchés    : {len(MARCHES)}")
print(f"  Sites      : {grand_total_sites}")
print(f"  Équipements: {grand_total_equips}")
print(f"{'='*70}")
print("\n✅ Preview terminée. Aucune donnée insérée en BDD.")
print("   Lancez 'python seed.py' pour insérer les données.\n")
