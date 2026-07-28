from docx import Document

def prepare_template():
    doc = Document("C:/Users/hanae/Desktop/Stage PFA 2026/backend/app/templates/template_ancfcc.docx")
    
    # Table 0: Date
    doc.tables[0].rows[1].cells[1].text = "{{ date_intervention }}"
    
    # Table 1: Signatures
    doc.tables[1].rows[1].cells[0].text = "{{ nom_technicien }}"
    doc.tables[1].rows[1].cells[1].text = "{{ nom_responsable }}"
    
    # Table 2: Equipment details
    doc.tables[2].rows[0].cells[3].text = "{{ puissance_kva }}"
    doc.tables[2].rows[1].cells[1].text = "{{ zone }}"
    doc.tables[2].rows[1].cells[3].text = "{{ nb_batteries }}"
    doc.tables[2].rows[2].cells[1].text = "{{ ville }}"
    doc.tables[2].rows[2].cells[3].text = "{{ marque_modele }}"
    doc.tables[2].rows[3].cells[1].text = "{{ etablissement }}"
    doc.tables[2].rows[3].cells[3].text = "{{ nom_site }}"
    doc.tables[2].rows[4].cells[3].text = "{{ numero_serie }}"
    doc.tables[2].rows[5].cells[3].text = "{{ capacite_batteries }}"
    
    # Table 3: Checklist
    # Rows 2 to 11 are the 10 points
    for i in range(2, 12):
        doc.tables[3].rows[i].cells[1].text = f"{{{{ pt{i-1}_oui }}}}"
        doc.tables[3].rows[i].cells[2].text = f"{{{{ pt{i-1}_obs }}}}"
        
    doc.save("C:/Users/hanae/Desktop/Stage PFA 2026/backend/app/templates/template_ancfcc.docx")
    print("Template prepared successfully.")

if __name__ == "__main__":
    prepare_template()
