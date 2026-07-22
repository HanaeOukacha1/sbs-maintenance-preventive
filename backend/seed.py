"""
========================================================
SEED COMPLET SBS — Version 2.0 (Lecture intégrale)
========================================================
Lit TOUS les fichiers (Excel ET Word), TOUTES les lignes,
détecte la structure automatiquement par marché.

Commande : python seed.py
(depuis backend/, avec venv actif)
========================================================
"""
import sys, os, re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from collections import defaultdict
from sqlalchemy.orm import Session
from app.db.database import SessionLocal
import app.models  # noqa

from app.models.marche import Marche
from app.models.site import Site
from app.models.equipement import Equipement, TypeEquipementEnum

try:
    import openpyxl
    import xlrd
    from docx import Document
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "openpyxl", "xlrd", "python-docx", "-q"])
    import openpyxl, xlrd
    from docx import Document

MASTER_DATA = os.path.normpath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "MASTER DATA")
)

# ============================================================
# MAPPING ARTICLE → TYPE ENUM
# ============================================================
def mapper_type(article: str) -> TypeEquipementEnum:
    a = str(article).upper().strip()
    if any(x in a for x in ["PORTABLE", "LAPTOP", "NOTEBOOK", " PT ", "PC PORT"]):
        return TypeEquipementEnum.PORTABLE
    if any(x in a for x in ["UC", "UNITE CENTRALE", "DESKTOP", "PC BURE", "ORDINATEUR DE BUREAU", "POSTE DE TRAVAIL"]):
        return TypeEquipementEnum.PC
    if a in ["PC"] or a.startswith("PC "):
        return TypeEquipementEnum.PC
    if any(x in a for x in ["SERVEUR", "SERVER", "RACK"]):
        return TypeEquipementEnum.SERVEUR
    if any(x in a for x in ["ONDULEUR", "UPS", "ALIMENTATION SANS COUPURE", "ASC"]):
        return TypeEquipementEnum.ONDULEUR
    if any(x in a for x in ["BAIE", "SWITCH", "ROUTEUR", "FIREWALL", "BRASSAGE", "STOCKAGE", "NAS", "SAN"]):
        return TypeEquipementEnum.BAIE_BRASSAGE
    if any(x in a for x in ["IMPRIMANTE", "IMP ", "MFP", "LASER", "JET D", "COPIEUR", "MULTIF", "FAX"]):
        return TypeEquipementEnum.IMPRIMANTE
    if any(x in a for x in ["ECRAN", "ÉCRAN", "MONITOR", "MONITEUR", "AFFICHEUR"]):
        return TypeEquipementEnum.ECRAN
    if any(x in a for x in ["SCANNER", "SCAN "]):
        return TypeEquipementEnum.SCANNER
    return TypeEquipementEnum.AUTRE

# ============================================================
# LECTURE EXCEL — TOUTES LES LIGNES
# ============================================================
def lire_xlsx_complet(filepath: str) -> dict:
    result = {}
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        for sn in wb.sheetnames:
            ws = wb[sn]
            rows = []
            for row in ws.iter_rows(values_only=True):
                r = [("" if c is None else str(c).strip()) for c in row]
                if any(r):
                    rows.append(r)
            result[sn] = rows
        wb.close()
    except Exception as e:
        print(f"    ⚠️  xlsx {filepath}: {e}")
    return result

def lire_xls_complet(filepath: str) -> dict:
    result = {}
    try:
        wb = xlrd.open_workbook(filepath)
        for sn in wb.sheet_names():
            ws = wb.sheet_by_name(sn)
            rows = []
            for i in range(ws.nrows):
                r = [str(ws.cell_value(i, j)).strip() for j in range(ws.ncols)]
                if any(r):
                    rows.append(r)
            result[sn] = rows
    except Exception as e:
        print(f"    ⚠️  xls {filepath}: {e}")
    return result

def lire_excel(filepath: str) -> dict:
    ext = filepath.lower().rsplit(".", 1)[-1]
    return lire_xlsx_complet(filepath) if ext == "xlsx" else lire_xls_complet(filepath)

# ============================================================
# LECTURE WORD — TOUTES LES LIGNES DE TOUS LES TABLEAUX
# ============================================================
def lire_docx_equipements(filepath: str) -> list:
    """
    Lit un fichier Word et extrait les équipements des tableaux.
    Retourne liste de dicts {site, type, marque, modele, serie, description}
    """
    equipements = []
    try:
        doc = Document(filepath)
        site_courant = ""
        # Chercher le site dans les textes du document
        for para in doc.paragraphs:
            t = para.text.strip().upper()
            if "VILLE" in t or "SITE :" in t:
                # Extraire le nom de ville
                for mot in t.split():
                    if len(mot) > 3 and mot.isalpha():
                        site_courant = mot
                        break

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = []
                prev = None
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt != prev:
                        cells.append(txt)
                    prev = txt
                if any(cells):
                    rows.append(cells)

            if not rows:
                continue

            # Chercher l'en-tête avec Désignation/N° Série
            header_idx = -1
            col_type = col_serie = col_marque = col_modele = col_obs = -1
            for i, row in enumerate(rows):
                row_u = [c.upper() for c in row]
                for j, c in enumerate(row_u):
                    if any(x in c for x in ["DÉSIGNATION", "DESIGNATION", "ARTICLE", "TYPE"]):
                        col_type = j
                    if any(x in c for x in ["N° SÉRIE", "N° SERIE", "SÉRIE", "SERIE", "SN"]):
                        col_serie = j
                    if "MARQUE" in c:
                        col_marque = j
                    if "MODÈLE" in c or "MODELE" in c:
                        col_modele = j
                    if any(x in c for x in ["OBSERVATION", "OBS", "ÉTAT", "ETAT"]):
                        col_obs = j
                if col_type >= 0 or col_serie >= 0:
                    header_idx = i
                    break

            if header_idx < 0:
                continue

            for row in rows[header_idx + 1:]:
                if not any(row):
                    continue
                type_val = row[col_type].strip() if col_type >= 0 and col_type < len(row) else ""
                serie_val = row[col_serie].strip() if col_serie >= 0 and col_serie < len(row) else ""
                marque_val = row[col_marque].strip() if col_marque >= 0 and col_marque < len(row) else ""
                modele_val = row[col_modele].strip() if col_modele >= 0 and col_modele < len(row) else ""
                obs_val = row[col_obs].strip() if col_obs >= 0 and col_obs < len(row) else ""

                if not type_val or type_val.upper() in ["DÉSIGNATION", "DESIGNATION", "TYPE", "ARTICLE"]:
                    continue
                if not any([marque_val, modele_val, serie_val, type_val]):
                    continue

                equipements.append({
                    "site": site_courant,
                    "type": mapper_type(type_val),
                    "nom": f"{type_val} {marque_val} {modele_val}".strip()[:200],
                    "marque": marque_val or None,
                    "modele": modele_val or None,
                    "serie": serie_val if serie_val and serie_val not in ["", "…", "...."] else None,
                    "description": obs_val or None,
                })
    except Exception as e:
        print(f"    ⚠️  docx {filepath}: {e}")
    return equipements

# ============================================================
# EXTRACTION INTELLIGENTE DEPUIS FEUILLE EXCEL
# ============================================================
def extraire_depuis_feuille(rows: list) -> dict:
    """
    Analyse toutes les lignes d'une feuille Excel.
    Retourne: {site_nom: [équipements]}
    Si pas de colonne SITE → site_nom = ""
    """
    # 1. Trouver la ligne d'en-tête
    header_idx = -1
    col = {}
    for i, row in enumerate(rows):
        r_up = [str(c).upper().strip() for c in row]
        found = 0
        tmp = {}
        for j, cell in enumerate(r_up):
            if cell in ["N° DE SERIE", "N° SERIE", "SN", "NUMÉRO DE SÉRIE", "NUMERO DE SERIE", "N°SERIE"]:
                tmp["serie"] = j; found += 1
            elif cell in ["MARQUE", "BRAND", "FABRICANT"]:
                tmp["marque"] = j; found += 1
            elif cell in ["MODELE", "MODÈLE", "MODEL", "MODÉLE"]:
                tmp["modele"] = j; found += 1
            elif cell in ["ARTICLE", "TYPE", "DÉSIGNATION", "DESIGNATION", "EQUIPEMENT", "MATÉRIEL", "MATERIEL"]:
                tmp["type"] = j; found += 1
            elif cell in ["AFFECTATION", "UTILISATEUR", "USER", "NOM UTILISATEUR"]:
                tmp["utilisateur"] = j
            elif cell in ["OBSERVATION", "OBSERVATIONS", "STATUT", "ETAT", "ÉTAT", "OBS"]:
                tmp["observation"] = j
            elif cell in ["SITE", "SITES", "LOCALISATION", "VILLE"]:
                tmp["site"] = j
            elif cell in ["N INVENTAIRE", "N° INVENTAIRE", "NO INVENTAIRE", "INVENTAIRE", "N°INV"]:
                tmp["inventaire"] = j
            elif cell in ["CPU", "PROCESSEUR"]:
                tmp["cpu"] = j
            elif cell in ["RAM", "MÉMOIRE", "MEMOIRE"]:
                tmp["ram"] = j
            elif cell in ["EMPLACEMENT", "LOCALISATION"]:
                tmp["emplacement"] = j

        if found >= 2:
            header_idx = i
            col = tmp
            break

    if header_idx < 0 or not col:
        return {}

    # 2. Lire les données
    par_site = defaultdict(list)
    has_site_col = "site" in col

    for row in rows[header_idx + 1:]:
        if not any(str(c).strip() for c in row):
            continue

        def get(key, row=row):
            idx = col.get(key)
            return str(row[idx]).strip() if idx is not None and idx < len(row) else ""

        type_val = get("type")
        marque_val = get("marque")
        modele_val = get("modele")
        serie_val = get("serie")
        site_val = get("site") if has_site_col else ""
        utilisateur_val = get("utilisateur")
        obs_val = get("observation")

        # Filtrer les lignes vides / sous-totaux
        if not type_val or type_val.upper() in [
            "", "TYPE", "ARTICLE", "DÉSIGNATION", "DESIGNATION",
            "TOTAL", "SOUS-TOTAL", "SOUS TOTAL", "NAN", "0.0", "0"
        ]:
            continue
        if not any([marque_val, modele_val, serie_val]):
            continue

        # Nettoyer N° série
        if serie_val.endswith(".0"):
            serie_val = serie_val[:-2]
        if serie_val.upper() in ["NAN", "N/A", "NONE", "", "0.0", "0"]:
            serie_val = None

        nom = f"{type_val} {marque_val} {modele_val}".strip()
        if utilisateur_val:
            nom = f"{nom} ({utilisateur_val})"

        par_site[site_val].append({
            "nom": nom[:200],
            "type": mapper_type(type_val),
            "marque": marque_val[:100] if marque_val else None,
            "modele": modele_val[:150] if modele_val else None,
            "serie": serie_val[:150] if serie_val else None,
            "description": obs_val[:500] if obs_val else None,
        })

    return dict(par_site)

# ============================================================
# NOM DE SITE DEPUIS LE NOM DE FICHIER
# ============================================================
def site_depuis_fichier(filename: str) -> str:
    nom = os.path.splitext(filename)[0]
    # Supprimer les préfixes
    for pfx in ["NV MP ", "MP ", "MD ", "Masters Data ", "Checklist ANCFCC 132 ", "Checklist ANCFCC 132"]:
        if nom.upper().startswith(pfx.upper()):
            nom = nom[len(pfx):]
    # Supprimer les suffixes de semestre/période
    nom = re.sub(r'\s+(S1|S2|1T|2T|3T|4T)[-_]?\d{0,4}.*$', '', nom, flags=re.IGNORECASE)
    nom = re.sub(r'\s+\d{4}$', '', nom)
    # Supprimer les marqueurs qualité
    for tag in [" ok", " OK", " MODIFIER", " RESEAUX", " MODIFIER OK", " PS "]:
        nom = nom.replace(tag, "")
    return nom.strip()

# ============================================================
# TRAITEMENT PAR MARCHÉ
# ============================================================
def traiter_marche(code: str, config: dict) -> dict:
    """
    Retourne {nom_site: [équipements]}
    """
    par_site = defaultdict(list)
    dossier = os.path.join(MASTER_DATA, code)

    if not os.path.isdir(dossier):
        return {}

    fichiers = sorted(os.listdir(dossier))

    for fichier in fichiers:
        path = os.path.join(dossier, fichier)
        ext = fichier.lower().rsplit(".", 1)[-1]

        # ---- WORD ----
        if ext == "docx":
            equips = lire_docx_equipements(path)
            # Pour ANCFCC : le site = ville dans le tableau Word
            if code == "ANCFCC":
                # Extraire la ville depuis le tableau 2 du document
                try:
                    doc = Document(path)
                    ville = ""
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [c.text.strip() for c in row.cells]
                            if "Ville" in cells or any("VILLE" == c.upper() for c in cells):
                                idx = next((i for i, c in enumerate(cells) if c.upper() == "VILLE"), -1)
                                if idx >= 0 and idx + 1 < len(cells):
                                    ville = cells[idx + 1].strip()
                            # Chercher aussi la valeur à côté de "Ville"
                            for i, c in enumerate(cells):
                                if c.upper() == "VILLE" and i + 1 < len(cells):
                                    ville = cells[i + 1].strip()
                    # Fallback : extraire depuis le nom de fichier
                    if not ville:
                        ville = site_depuis_fichier(fichier)

                    # Ajouter l'onduleur depuis le tableau (N° Série trouvé dans tableau 2)
                    for table in doc.tables:
                        for row in table.rows:
                            cells_u = [c.text.strip().upper() for c in row.cells]
                            if "N° SÉRIE" in cells_u or "N° SERIE" in cells_u:
                                # Lire la ligne suivante
                                pass
                    # Récupérer les données de l'onduleur
                    serie_ond = ""
                    marque_ond = ""
                    for table in doc.tables:
                        rows_t = [[c.text.strip() for c in row.cells] for row in table.rows]
                        for i, r in enumerate(rows_t):
                            if any("N° SÉRIE" in c.upper() or "N° SERIE" in c.upper() for c in r):
                                # Chercher la valeur dans la même ligne ou la suivante
                                for j, c in enumerate(r):
                                    if "N° SÉRIE" in c.upper() or "N° SERIE" in c.upper():
                                        if j + 1 < len(r) and r[j + 1]:
                                            serie_ond = r[j + 1]
                            if any("MARQUE" in c.upper() for c in r):
                                for j, c in enumerate(r):
                                    if "MARQUE" in c.upper() and j + 1 < len(r):
                                        marque_ond = r[j + 1]

                    if ville:
                        par_site[ville].append({
                            "nom": f"Onduleur {marque_ond}".strip() or "Onduleur",
                            "type": TypeEquipementEnum.ONDULEUR,
                            "marque": marque_ond or None,
                            "modele": None,
                            "serie": serie_ond or None,
                            "description": f"Site ANCFCC {ville}",
                        })
                except Exception as e:
                    print(f"    ⚠️  ANCFCC docx: {e}")

            elif code == "ADM":
                # ADM Word : tableau de serveurs
                try:
                    doc = Document(path)
                    for table in doc.tables:
                        rows_t = [[c.text.strip() for c in row.cells] for row in table.rows]
                        header_found = False
                        col_des = col_marq = col_mod = col_ser = col_cpu = -1
                        for i, r in enumerate(rows_t):
                            r_up = [c.upper() for c in r]
                            if "DÉSIGNATION" in r_up or "DESIGNATION" in r_up:
                                for j, c in enumerate(r_up):
                                    if "DÉSIGNATION" in c or "DESIGNATION" in c: col_des = j
                                    if "FABRICANT" in c or "MARQUE" in c: col_marq = j
                                    if "MODÈLE" in c or "MODELE" in c: col_mod = j
                                    if "SÉRIE" in c or "SERIE" in c: col_ser = j
                                    if "PROCESSEUR" in c or "CPU" in c: col_cpu = j
                                header_found = True
                                continue
                            if header_found:
                                des = r[col_des] if col_des >= 0 and col_des < len(r) else ""
                                marq = r[col_marq] if col_marq >= 0 and col_marq < len(r) else ""
                                mod = r[col_mod] if col_mod >= 0 and col_mod < len(r) else ""
                                ser = r[col_ser] if col_ser >= 0 and col_ser < len(r) else ""
                                cpu = r[col_cpu] if col_cpu >= 0 and col_cpu < len(r) else ""
                                if des and des.upper() not in ["", "DÉSIGNATION", "DESIGNATION"]:
                                    par_site["Siège ADM"].append({
                                        "nom": f"{des} {marq} {mod}".strip()[:200],
                                        "type": TypeEquipementEnum.SERVEUR,
                                        "marque": marq or None,
                                        "modele": mod or None,
                                        "serie": ser or None,
                                        "description": f"CPU: {cpu}" if cpu else None,
                                    })
                except Exception as e:
                    print(f"    ⚠️  ADM docx: {e}")

            elif code == "MHAI":
                # Word MHAI → on ignore pour le seed équipements (ils sont dans l'Excel)
                pass

            continue

        # ---- EXCEL / XLS ----
        if ext not in ("xlsx", "xls"):
            continue

        print(f"    📄 {fichier}")
        feuilles = lire_excel(path)

        # ---- Logique selon le marché ----

        if code in ("CNDH", "MSANTE", "AMEE"):
            # Fichier = 1 site, colonnes dans les feuilles
            nom_site = site_depuis_fichier(fichier)
            for nom_feuille, rows in feuilles.items():
                result = extraire_depuis_feuille(rows)
                for site_col, equips in result.items():
                    # Si la colonne SITE existe dans les données, utiliser ça
                    s = site_col if site_col else nom_site
                    par_site[s].extend(equips)

        elif code in ("ANP",):
            # Feuille = 1 site (nom de la feuille = site)
            for nom_feuille, rows in feuilles.items():
                if nom_feuille.lower() in ["feuil1", "feuil2", "feuil3", "sheet1", "sheet2"]:
                    # Ces feuilles peuvent aussi avoir des données utiles
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        s = site_col if site_col else nom_feuille
                        par_site[s].extend(equips)
                else:
                    result = extraire_depuis_feuille(rows)
                    for site_col, equips in result.items():
                        s = site_col if site_col else nom_feuille
                        par_site[s].extend(equips)

        elif code in ("ONP", "MARSA MAROC", "AOH", "INPPLC", "MHAI"):
            # Peut avoir colonne SITE ou feuille par site
            for nom_feuille, rows in feuilles.items():
                result = extraire_depuis_feuille(rows)
                for site_col, equips in result.items():
                    if site_col:
                        # Colonne SITE présente → grouper par site
                        par_site[site_col].extend(equips)
                    elif nom_feuille.lower() not in ["feuil1", "feuil2", "feuil3", "sheet1"]:
                        # Nom de feuille = site
                        par_site[nom_feuille.strip()].extend(equips)
                    else:
                        # Fallback : nom du marché
                        par_site[f"Siège {code}"].extend(equips)

        else:
            # Générique
            for nom_feuille, rows in feuilles.items():
                result = extraire_depuis_feuille(rows)
                for site_col, equips in result.items():
                    s = site_col if site_col else nom_feuille
                    par_site[s].extend(equips)

    return dict(par_site)

# ============================================================
# MARCHÉS STATIQUES (sites sans fichiers exploitables)
# ============================================================
MARCHES = {
    "ADM":         "Administration",
    "AMEE":        "Agence Marocaine pour l'Efficacité Énergétique",
    "ANCFCC":      "Agence Nationale de la Conservation Foncière, du Cadastre et de la Cartographie",
    "ANP":         "Agence Nationale des Ports",
    "AOH":         "Al Omrane Holding",
    "CNDH":        "Conseil National des Droits de l'Homme",
    "INPPLC":      "Instance Nationale de la Probité, de la Prévention et de la Lutte contre la Corruption",
    "MARSA MAROC": "Marsa Maroc",
    "MHAI":        "Ministère des Habous et des Affaires Islamiques",
    "MSANTE":      "Ministère de la Santé",
    "ONP":         "Office National des Pêches",
}

# Sites ANCFCC (extraits des noms de fichiers Word)
SITES_ANCFCC = [
    "AGADIR", "BENGUERIR", "KHEMISSAT", "KENITRA", "MDIEQ",
    "MIDELT", "OUARZAZATE", "ROMMANI", "SALA EL JADIDA", "SETTAT",
    "SIDI SLIMANE", "TETOUAN", "SIDI BANNOUR"
]

# ============================================================
# SEED PRINCIPAL
# ============================================================
def seed():
    db: Session = SessionLocal()
    total_m = total_s = total_e = 0

    try:
        print("\n" + "="*60)
        print("SEED COMPLET SBS v2.0")
        print("="*60)

        for code, client in MARCHES.items():
            print(f"\n📁 {code} — {client}")

            # Créer/récupérer le marché
            marche = db.query(Marche).filter(Marche.nom == code).first()
            if not marche:
                marche = Marche(
                    nom=code, client=client,
                    description=f"Maintenance préventive IT — {client}",
                    is_active=True
                )
                db.add(marche); db.flush()
                total_m += 1
                print(f"  ✅ Marché créé (id={marche.id})")
            else:
                print(f"  ℹ️  Existant (id={marche.id})")

            # Extraire toutes les données des fichiers
            par_site = traiter_marche(code, {})

            # Pour ANCFCC : ajouter les 13 sites même si Word n'a pas d'équipements parsables
            if code == "ANCFCC":
                for ville in SITES_ANCFCC:
                    if ville not in par_site:
                        par_site[ville] = []

            # Si aucun site trouvé → créer un site par défaut
            if not par_site:
                par_site[f"Siège {code}"] = []

            # Créer sites et équipements
            for nom_site, equips in sorted(par_site.items()):
                if not nom_site.strip():
                    nom_site = f"Siège {code}"

                site = db.query(Site).filter(
                    Site.nom == nom_site, Site.marche_id == marche.id
                ).first()
                if not site:
                    site = Site(nom=nom_site, marche_id=marche.id, is_active=True)
                    db.add(site); db.flush()
                    total_s += 1

                print(f"  📍 {nom_site} → {len(equips)} équipement(s)")

                for eq in equips:
                    # Dédupliquer par (serie, site)
                    if eq.get("serie"):
                        exists = db.query(Equipement).filter(
                            Equipement.numero_serie == eq["serie"],
                            Equipement.site_id == site.id
                        ).first()
                        if exists:
                            continue

                    e = Equipement(
                        nom=eq["nom"],
                        type_equipement=eq["type"],
                        marque=eq.get("marque"),
                        modele=eq.get("modele"),
                        numero_serie=eq.get("serie"),
                        description=eq.get("description"),
                        site_id=site.id,
                        is_active=True,
                    )
                    db.add(e)
                    total_e += 1

        db.commit()
        print("\n" + "="*60)
        print("✅ SEED TERMINÉ !")
        print(f"   Marchés  : {total_m}")
        print(f"   Sites    : {total_s}")
        print(f"   Équip.   : {total_e}")
        print("="*60)

    except Exception as ex:
        db.rollback()
        print(f"\n❌ ERREUR: {ex}")
        import traceback; traceback.print_exc()
    finally:
        db.close()

if __name__ == "__main__":
    seed()
