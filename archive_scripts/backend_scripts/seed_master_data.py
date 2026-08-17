"""
Script de seeding — Import complet des Master Data SBS vers MySQL.

Usage (depuis le dossier backend, avec venv activé) :
    python seed_master_data.py

Ce script :
1. Crée les marchés
2. Crée les sites avec leur checklist_type et feuilles
3. Importe tous les équipements depuis les fichiers Excel/Word
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.db.database import SessionLocal
from app.models import (
    Marche, Site, Equipement,
    TypeEquipementEnum, ChecklistTypeEnum
)

# Chemin vers le dossier MASTER DATA
MASTER_DATA_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'MASTER DATA')


def get_type(designation: str) -> TypeEquipementEnum:
    """Détermine le TypeEquipementEnum depuis une désignation texte."""
    d = str(designation).upper().strip()
    if any(k in d for k in ['SERVEUR', 'SERVER', 'BAIE DE']):
        return TypeEquipementEnum.SERVEUR
    if 'KVM' in d:
        return TypeEquipementEnum.KVM
    if any(k in d for k in ['AIO', 'ALL-IN-ONE', 'ALL IN ONE']):
        return TypeEquipementEnum.AIO
    if any(k in d for k in ['PORTABLE', 'LAPTOP', 'NOTEBOOK']):
        return TypeEquipementEnum.PORTABLE
    if any(k in d for k in ['UC', 'FIXE', 'PC', 'UNITE CENTRALE', 'UNITÉ CENTRALE', 'BUREAU', 'DESKTOP']):
        return TypeEquipementEnum.PC
    if any(k in d for k in ['IMP', 'MFP', 'LASER', 'MULTIFONCTION', 'FAX']):
        if 'FAX' in d and 'IMP' not in d and 'MFP' not in d:
            return TypeEquipementEnum.FAX
        return TypeEquipementEnum.IMPRIMANTE
    if 'FAX' in d:
        return TypeEquipementEnum.FAX
    if any(k in d for k in ['ECRAN', 'ÉCRAN', 'MONITOR', 'MONITEUR', 'SCREEN']):
        return TypeEquipementEnum.ECRAN
    if 'SCANNER' in d or 'SCAN' in d:
        return TypeEquipementEnum.SCANNER
    if 'ONDULEUR' in d or 'UPS' in d:
        return TypeEquipementEnum.ONDULEUR
    if 'PHOTO' in d:
        return TypeEquipementEnum.PHOTOCOPIEUR
    return TypeEquipementEnum.AUTRE


def clean(val):
    """Nettoie une valeur (None, espaces, etc.)"""
    if val is None:
        return None
    s = str(val).strip()
    return s if s and s.lower() not in ['none', 'nan', ''] else None


def seed_adm(db, marche_id):
    """ADM — 1 site, serveurs principaux + redondants"""
    import docx
    site = Site(
        nom="ADM Siège",
        ville="Rabat",
        marche_id=marche_id,
        checklist_type=ChecklistTypeEnum.ADM,
        feuilles=None
    )
    db.add(site)
    db.flush()

    doc_path = os.path.join(MASTER_DATA_PATH, 'ADM', 'MP ADM.docx')
    if not os.path.exists(doc_path):
        print(f"  ⚠️  Fichier ADM non trouvé : {doc_path}")
        return

    doc = docx.Document(doc_path)
    principal_id = None

    for table in doc.tables:
        for row in table.rows[1:]:  # Skip header
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            designation = clean(cells[0])
            if not designation:
                continue

            marque = clean(cells[1])
            modele = clean(cells[2])
            processeur = clean(cells[4]) if len(cells) > 4 else None
            memoire = clean(cells[5]) if len(cells) > 5 else None
            disque_c = clean(cells[6]) if len(cells) > 6 else None
            disque_d = clean(cells[7]) if len(cells) > 7 else None
            ip = clean(cells[8]) if len(cells) > 8 else None

            is_redondant = 'REDONDANT' in designation.upper()
            equip = Equipement(
                nom=f"{marque or ''} {modele or ''}".strip() or designation,
                designation=designation,
                type_equipement=TypeEquipementEnum.SERVEUR,
                site_id=site.id,
                marque=marque,
                modele=modele,
                cpu=processeur,
                ram=memoire,
                disque_dur=f"C:{disque_c} D:{disque_d}" if disque_c else None,
                ip=ip,
                est_serveur_redondant=is_redondant,
                serveur_principal_id=principal_id if is_redondant else None,
            )
            db.add(equip)
            db.flush()
            if not is_redondant:
                principal_id = equip.id

    print(f"  ✅ ADM : équipements importés")


def seed_amee(db, marche_id):
    """AMEE — Marrakech (2 feuilles) + Rabat (4 feuilles)"""
    import openpyxl

    sites_config = [
        {
            "nom": "AMEE Marrakech",
            "ville": "Marrakech",
            "checklist_type": ChecklistTypeEnum.AMEE_MARRAKECH,
            "feuilles": ["Imprimantes & MFP", "Serveurs"],
            "file": "MP AMEE MARRAKECH 1T-2026.xlsx",
        },
        {
            "nom": "AMEE Rabat",
            "ville": "Rabat",
            "checklist_type": ChecklistTypeEnum.AMEE_RABAT,
            "feuilles": ["PC", "MàJ Windows", "Imp & MFP Réseaux", "Data Center"],
            "file": "NV MP AMEE RABAT 1T-2026.xlsx",
        },
    ]

    for config in sites_config:
        site = Site(
            nom=config["nom"],
            ville=config["ville"],
            marche_id=marche_id,
            checklist_type=config["checklist_type"],
            feuilles=config["feuilles"]
        )
        db.add(site)
        db.flush()

        file_path = os.path.join(MASTER_DATA_PATH, 'AMEE', config["file"])
        if not os.path.exists(file_path):
            print(f"  ⚠️  Fichier AMEE non trouvé : {file_path}")
            continue

        wb = openpyxl.load_workbook(file_path)
        count = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers = []
            header_row = None

            # Trouver la ligne d'en-tête
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                row_str = [str(c).lower() if c else '' for c in row]
                if any(k in ' '.join(row_str) for k in ['série', 'serie', 's/n', 'désignation', 'designation', 'type', 'utilisateur']):
                    headers = [str(c).strip() if c else '' for c in row]
                    header_row = i
                    break

            if not headers:
                continue

            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i <= header_row:
                    continue
                if not any(c for c in row):
                    continue

                data = dict(zip(headers, row))

                # Récupérer N° de série (chercher dans plusieurs colonnes possibles)
                nserie = None
                for key in data:
                    if 'serie' in key.lower() or 's/n' in key.lower():
                        nserie = clean(data[key])
                        break

                if not nserie:
                    continue

                # Désignation / Type
                designation = None
                for key in data:
                    if any(k in key.lower() for k in ['désig', 'desig', 'type', 'famil']):
                        designation = clean(data[key])
                        break

                marque = None
                for key in data:
                    if 'marque' in key.lower() or 'brand' in key.lower():
                        marque = clean(data[key])
                        break

                modele = None
                for key in data:
                    if any(k in key.lower() for k in ['modèle', 'modele', 'model', 'désign', 'article']):
                        if key not in ['Désignation', 'Designation']:
                            modele = clean(data[key])
                            break

                utilisateur = None
                for key in data:
                    if any(k in key.lower() for k in ['personne', 'utilisateur', 'nom']):
                        utilisateur = clean(data[key])
                        break

                cpu = None
                for key in data:
                    if 'cpu' in key.lower() or 'processeur' in key.lower():
                        cpu = clean(data[key])
                        break

                ram = None
                for key in data:
                    if 'ram' in key.lower() or 'mémoire' in key.lower() or 'memoire' in key.lower():
                        ram = clean(data[key])
                        break

                os_val = None
                for key in data:
                    if 'système' in key.lower() or 'systeme' in key.lower() or 'os' in key.lower():
                        os_val = clean(data[key])
                        break

                equip = Equipement(
                    nom=f"{marque or ''} {modele or ''}".strip() or designation or 'Équipement',
                    designation=designation,
                    type_equipement=get_type(designation or ''),
                    site_id=site.id,
                    marque=marque,
                    modele=modele,
                    numero_serie=nserie,
                    utilisateur_nom=utilisateur,
                    cpu=cpu,
                    ram=ram,
                    systeme_exploitation=os_val,
                    sous_site=sheet_name,  # nom de la feuille Excel comme sous-site
                )
                db.add(equip)
                count += 1

        print(f"  ✅ {config['nom']} : {count} équipements importés")


def seed_ancfcc(db, marche_id):
    """ANCFCC — 13 sites, chacun avec 1 onduleur"""
    import docx

    sites_ancfcc = [
        ("Agadir", "SUD", "CADASTRE", "Riello 15 KVA", "LU53UT895900002"),
        ("Benguerir", "SUD", "CADASTRE", None, None),
        ("Khemissat", "CENTRE", "CADASTRE", None, None),
        ("Kenitra", "NORD", "CADASTRE", None, None),
        ("M'dieq", "NORD", "CADASTRE", None, None),
        ("Midelt", "CENTRE", "CADASTRE", None, None),
        ("Ouarzazate", "SUD", "CADASTRE", None, None),
        ("Rommani", "CENTRE", "CADASTRE", None, None),
        ("Salé El Jadida", "CENTRE", "CADASTRE", None, None),
        ("Settat", "CENTRE", "CADASTRE", None, None),
        ("Sidi Slimane", "NORD", "CADASTRE", None, None),
        ("Tétouan", "NORD", "CADASTRE", None, None),
        ("Sidi Bennour", "SUD", "CADASTRE", None, None),
    ]

    # Lire les N° de série depuis les fichiers Word si disponibles
    ancfcc_path = os.path.join(MASTER_DATA_PATH, 'ANCFCC')
    nserie_map = {}
    if os.path.exists(ancfcc_path):
        for fname in os.listdir(ancfcc_path):
            if fname.endswith('.docx'):
                try:
                    doc = docx.Document(os.path.join(ancfcc_path, fname))
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [c.text.strip() for c in row.cells]
                            for i, c in enumerate(cells):
                                if 'série' in c.lower() or 'serie' in c.lower():
                                    if i + 1 < len(cells) and cells[i + 1]:
                                        # Extraire la ville depuis le nom du fichier
                                        ville_key = fname.replace('Checklist ANCFCC 132 ', '').replace('.docx', '').strip().upper()
                                        nserie_map[ville_key] = cells[i + 1]
                except Exception:
                    pass

    for ville, zone, etablissement, marque_modele, nserie_default in sites_ancfcc:
        site = Site(
            nom=f"ANCFCC {ville}",
            ville=ville,
            marche_id=marche_id,
            checklist_type=ChecklistTypeEnum.ANCFCC,
            feuilles=None
        )
        db.add(site)
        db.flush()

        # Chercher le N° de série dans la map
        nserie = nserie_map.get(ville.upper(), nserie_default)

        equip = Equipement(
            nom=marque_modele or "Onduleur Riello",
            designation="ONDULEUR",
            type_equipement=TypeEquipementEnum.ONDULEUR,
            site_id=site.id,
            marque="Riello" if marque_modele and 'riello' in marque_modele.lower() else None,
            modele=marque_modele,
            numero_serie=nserie,
            zone=zone,
            entite=etablissement,
        )
        db.add(equip)

    print(f"  ✅ ANCFCC : 13 sites + onduleurs importés")


def seed_excel_simple(db, marche_id, file_path, sites_config, checklist_type,
                      has_numero_inventaire=False, feuilles=None):
    """
    Import générique pour marchés simples (ANP, AOH, INPPLC, MHAI, MSANTE, ONP, MARSA, CNDH).
    """
    import openpyxl
    import xlrd

    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.xlsx':
        wb = openpyxl.load_workbook(file_path)
        sheets = wb.sheetnames
    elif ext in ['.xls']:
        wb_xls = xlrd.open_workbook(file_path)
        sheets = wb_xls.sheet_names()
    else:
        print(f"  ⚠️  Format non supporté : {file_path}")
        return

    for sh_name in sheets:
        if sh_name in ['Feuil1', 'Feuil2', 'Sheet1', 'Sheet2'] and len(sheets) > 1:
            # Skip feuilles vides génériques si d'autres existent
            pass

        if ext == '.xlsx':
            ws = wb[sh_name]
            rows_data = list(ws.iter_rows(values_only=True))
        else:
            ws_xls = wb_xls.sheet_by_name(sh_name)
            rows_data = [ws_xls.row_values(i) for i in range(ws_xls.nrows)]

        # Trouver config du site pour cette feuille
        site_config = sites_config.get(sh_name, sites_config.get('__default__'))
        if not site_config:
            continue

        # Créer ou récupérer le site
        site = site_config.get('__site_obj__')
        if not site:
            site = Site(
                nom=site_config['nom'],
                ville=site_config.get('ville', ''),
                marche_id=marche_id,
                checklist_type=checklist_type,
                feuilles=site_config.get('feuilles', feuilles),
            )
            db.add(site)
            db.flush()
            site_config['__site_obj__'] = site

        # Trouver la ligne d'en-tête
        header_idx = None
        headers = []
        for i, row in enumerate(rows_data):
            row_str = ' '.join([str(c).lower() for c in row if c])
            if any(k in row_str for k in ['série', 'serie', 's/n', 'désignation', 'designation', 'article', 'famille', 'matériel', 'materiel']):
                headers = [str(c).strip() if c else '' for c in row]
                header_idx = i
                break

        if header_idx is None:
            continue

        count = 0
        for row in rows_data[header_idx + 1:]:
            if not any(c for c in row):
                continue

            data = dict(zip(headers, row))

            # N° série
            nserie = None
            for key in headers:
                if 'série' in key.lower() or 'serie' in key.lower() or 's/n' in key.lower():
                    nserie = clean(data.get(key))
                    break
            if not nserie:
                continue

            # Désignation
            designation = None
            for key in headers:
                if any(k in key.lower() for k in ['désig', 'desig', 'article', 'famil', 'matér', 'mater']):
                    designation = clean(data.get(key))
                    break

            marque = None
            for key in headers:
                if 'marque' in key.lower():
                    marque = clean(data.get(key))
                    break

            modele = None
            for key in headers:
                if any(k in key.lower() for k in ['modèle', 'modele', 'model']):
                    modele = clean(data.get(key))
                    break

            num_inv = None
            if has_numero_inventaire:
                for key in headers:
                    if 'inv' in key.lower():
                        num_inv = clean(data.get(key))
                        break

            # Champs spéciaux
            entite = None
            for key in headers:
                if 'entit' in key.lower():
                    entite = clean(data.get(key))
                    break

            emplacement = None
            for key in headers:
                if 'emplace' in key.lower() or 'bureau' in key.lower():
                    emplacement = clean(data.get(key))
                    break

            affectation = None
            for key in headers:
                if 'affect' in key.lower() or 'personn' in key.lower():
                    affectation = clean(data.get(key))
                    break

            direction = None
            for key in headers:
                if 'direction' in key.lower():
                    direction = clean(data.get(key))
                    break

            utilisateur = None
            for key in headers:
                if 'nom' in key.lower() and 'prénom' in key.lower():
                    utilisateur = clean(data.get(key))
                    break
                elif 'utilisateur' in key.lower():
                    utilisateur = clean(data.get(key))
                    break

            sous_site_val = site_config.get('sous_site')
            if not sous_site_val and sh_name not in ['A', 'Feuil1', 'Sheet1']:
                sous_site_val = sh_name

            equip = Equipement(
                nom=f"{marque or ''} {modele or ''}".strip() or designation or 'Équipement',
                designation=designation,
                type_equipement=get_type(designation or ''),
                site_id=site.id,
                marque=marque,
                modele=modele,
                numero_serie=nserie,
                numero_inventaire=num_inv,
                entite=entite,
                emplacement=emplacement,
                affectation=affectation,
                direction=direction,
                utilisateur_nom=utilisateur,
                sous_site=sous_site_val,
            )
            db.add(equip)
            count += 1

        print(f"  ✅ {site_config['nom']} ({sh_name}) : {count} équipements")


def main():
    print("🚀 Démarrage du seeding Master Data SBS...\n")
    db = SessionLocal()

    try:
        # ======================================================
        # VÉRIFICATION : Éviter double import
        # ======================================================
        if db.query(Marche).count() > 0:
            print("⚠️  Des données existent déjà. Supprimez-les manuellement si vous voulez réimporter.")
            print("    SQL : DELETE FROM equipements; DELETE FROM sites; DELETE FROM marches;")
            return

        # ======================================================
        # 1. ADM — Autoroutes du Maroc
        # ======================================================
        print("📦 ADM — Autoroutes du Maroc")
        marche_adm = Marche(nom="ADM", client="Autoroutes du Maroc", description="Marché Autoroutes du Maroc")
        db.add(marche_adm)
        db.flush()
        seed_adm(db, marche_adm.id)

        # ======================================================
        # 2. AMEE — Agence Marocaine de l'Efficacité Énergétique
        # ======================================================
        print("\n📦 AMEE — Agence Marocaine de l'Efficacité Énergétique")
        marche_amee = Marche(nom="AMEE", client="Agence Marocaine de l'Efficacité Énergétique")
        db.add(marche_amee)
        db.flush()
        seed_amee(db, marche_amee.id)

        # ======================================================
        # 3. ANCFCC — Agence Nationale de la Conservation Foncière
        # ======================================================
        print("\n📦 ANCFCC — Agence Nationale de la Conservation Foncière")
        marche_ancfcc = Marche(nom="ANCFCC", client="Agence Nationale de la Conservation Foncière du Cadastre et de la Cartographie")
        db.add(marche_ancfcc)
        db.flush()
        seed_ancfcc(db, marche_ancfcc.id)

        # ======================================================
        # 4. ANP — Agence Nationale des Ports
        # ======================================================
        print("\n📦 ANP — Agence Nationale des Ports")
        marche_anp = Marche(nom="ANP", client="Agence Nationale des Ports")
        db.add(marche_anp)
        db.flush()
        anp_path = os.path.join(MASTER_DATA_PATH, 'ANP', 'MP ANP.xlsx')
        anp_sites = {
            "Port El Jadida": {"nom": "ANP El Jadida", "ville": "El Jadida"},
            "Port Jorf Lasfar": {"nom": "ANP Jorf Lasfar", "ville": "Jorf Lasfar"},
        }
        seed_excel_simple(db, marche_anp.id, anp_path, anp_sites, ChecklistTypeEnum.ANP)

        # ======================================================
        # 5. AOH — Al Omrane Holding
        # ======================================================
        print("\n📦 AOH — Al Omrane Holding")
        marche_aoh = Marche(nom="AOH", client="Al Omrane Holding")
        db.add(marche_aoh)
        db.flush()
        aoh_path = os.path.join(MASTER_DATA_PATH, 'AOH', 'MD HAO.xlsx')
        aoh_sites = {"AOH": {"nom": "AOH Siège", "ville": "Rabat"}, "__default__": {"nom": "AOH Siège", "ville": "Rabat"}}
        seed_excel_simple(db, marche_aoh.id, aoh_path, aoh_sites, ChecklistTypeEnum.AOH, has_numero_inventaire=True)

        # ======================================================
        # 6. INPPLC
        # ======================================================
        print("\n📦 INPPLC")
        marche_inpplc = Marche(nom="INPPLC", client="Instance Nationale de la Probité, de la Prévention et de la Lutte contre la Corruption")
        db.add(marche_inpplc)
        db.flush()
        inpplc_path = os.path.join(MASTER_DATA_PATH, 'INPPLC', 'Masters Data INPPLC.XLS')
        site_inpplc = Site(nom="INPPLC Rabat", ville="Rabat", marche_id=marche_inpplc.id,
                           checklist_type=ChecklistTypeEnum.INPPLC,
                           feuilles=["Imprimantes", "PC Portables"])
        db.add(site_inpplc)
        db.flush()
        inpplc_sites = {
            "Imprimantes": {"nom": "INPPLC Rabat", "ville": "Rabat", "sous_site": "Imprimantes", "__site_obj__": site_inpplc},
            "Pc portables": {"nom": "INPPLC Rabat", "ville": "Rabat", "sous_site": "PC Portables", "__site_obj__": site_inpplc},
        }
        seed_excel_simple(db, marche_inpplc.id, inpplc_path, inpplc_sites, ChecklistTypeEnum.INPPLC)

        # ======================================================
        # 7. MARSA MAROC
        # ======================================================
        print("\n📦 Marsa Maroc")
        marche_mm = Marche(nom="MARSA_MAROC", client="Marsa Maroc")
        db.add(marche_mm)
        db.flush()
        mm_path = os.path.join(MASTER_DATA_PATH, 'MARSA MAROC', 'MP MM .xlsx')
        mm_sites = {"__default__": {"nom": "Marsa Maroc Siège", "ville": "Casablanca"}, "Sheet1 (2)": {"nom": "Marsa Maroc Siège", "ville": "Casablanca"}}
        seed_excel_simple(db, marche_mm.id, mm_path, mm_sites, ChecklistTypeEnum.MARSA_MAROC)

        # ======================================================
        # 8. MHAI — Ministère des Habous
        # ======================================================
        print("\n📦 MHAI — Ministère des Habous et des Affaires Islamiques")
        marche_mhai = Marche(nom="MHAI", client="Ministère des Habous et des Affaires Islamiques")
        db.add(marche_mhai)
        db.flush()
        mhai_path = os.path.join(MASTER_DATA_PATH, 'MHAI', 'MP HABOUS S2-24.xlsx')
        mhai_villes = {
            "TANGER": "Tanger", "MARRAKECH": "Marrakech", "CASABLANCA": "Casablanca",
            "RABAT": "Rabat", "SALE": "Salé", "OUJDA,": "Oujda", "Feuil1": None
        }
        mhai_sites = {}
        for sh, ville in mhai_villes.items():
            if ville:
                mhai_sites[sh] = {"nom": f"MHAI {ville}", "ville": ville}
        seed_excel_simple(db, marche_mhai.id, mhai_path, mhai_sites, ChecklistTypeEnum.MHAI, has_numero_inventaire=True)

        # ======================================================
        # 9. MSANTE — Ministère de la Santé
        # ======================================================
        print("\n📦 MSANTE — Ministère de la Santé")
        marche_msante = Marche(nom="MSANTE", client="Ministère de la Santé")
        db.add(marche_msante)
        db.flush()

        msante_files = [
            ("CCM S2.xlsx", "CCM", ChecklistTypeEnum.MSANTE_STANDARD, None),
            ("DIVISION D'INFORMATION S2.xlsx", "Division d'Information", ChecklistTypeEnum.MSANTE_STANDARD, None),
            ("INVENTAIRE DHSA PS .xlsx", "DHSA", ChecklistTypeEnum.MSANTE_STANDARD, None),
            ("RAMED S2.xlsx", "RAMED (Poste Comptable)", ChecklistTypeEnum.MSANTE_STANDARD, None),
        ]
        for fname, site_name, ctype, feuilles_list in msante_files:
            fpath = os.path.join(MASTER_DATA_PATH, 'MSANTE', fname)
            if os.path.exists(fpath):
                sites_cfg = {"__default__": {"nom": f"MSANTE Rabat — {site_name}", "ville": "Rabat", "feuilles": feuilles_list}}
                # Gérer les noms de feuilles connus
                for sh_try in ['A', 'Feuil1', 'Sheet1', 'DHSA']:
                    sites_cfg[sh_try] = sites_cfg["__default__"]
                seed_excel_simple(db, marche_msante.id, fpath, sites_cfg, ctype)

        # ======================================================
        # 10. ONP — Office National des Pêches
        # ======================================================
        print("\n📦 ONP — Office National des Pêches")
        marche_onp = Marche(nom="ONP", client="Office National des Pêches")
        db.add(marche_onp)
        db.flush()

        import openpyxl
        onp_path = os.path.join(MASTER_DATA_PATH, 'ONP', 'MP ONP.xlsx')
        if os.path.exists(onp_path):
            wb_onp = openpyxl.load_workbook(onp_path)
            ws_onp = wb_onp.active
            rows_onp = list(ws_onp.iter_rows(values_only=True))

            # Trouver header
            h_idx = None
            headers_onp = []
            for i, row in enumerate(rows_onp):
                row_str = ' '.join([str(c).lower() for c in row if c])
                if 'serie' in row_str or 'désig' in row_str:
                    headers_onp = [str(c).strip() if c else '' for c in row]
                    h_idx = i
                    break

            # Regrouper par SITE
            site_cache = {}
            count_onp = 0
            if h_idx is not None:
                for row in rows_onp[h_idx + 1:]:
                    if not any(c for c in row):
                        continue
                    data = dict(zip(headers_onp, row))
                    ville_onp = clean(data.get('SITE') or data.get('Site') or data.get('site'))
                    nserie_onp = None
                    for key in headers_onp:
                        if 'serie' in key.lower() or 's/n' in key.lower():
                            nserie_onp = clean(data.get(key))
                            break
                    if not ville_onp or not nserie_onp:
                        continue

                    if ville_onp not in site_cache:
                        s = Site(nom=f"ONP {ville_onp}", ville=ville_onp, marche_id=marche_onp.id,
                                 checklist_type=ChecklistTypeEnum.ONP)
                        db.add(s)
                        db.flush()
                        site_cache[ville_onp] = s

                    designation_onp = clean(data.get('DESIGNATION') or data.get('Désignation') or data.get('designation'))
                    marque_onp = clean(data.get('MARQUE') or data.get('Marque'))
                    modele_onp = clean(data.get('MODELE') or data.get('Modèle') or data.get('MODÈLE'))

                    equip = Equipement(
                        nom=f"{marque_onp or ''} {modele_onp or ''}".strip() or designation_onp or 'Équipement',
                        designation=designation_onp,
                        type_equipement=get_type(designation_onp or ''),
                        site_id=site_cache[ville_onp].id,
                        marque=marque_onp,
                        modele=modele_onp,
                        numero_serie=nserie_onp,
                    )
                    db.add(equip)
                    count_onp += 1

            print(f"  ✅ ONP : {count_onp} équipements dans {len(site_cache)} villes")

        # ======================================================
        # 11. CNDH — Conseil National des Droits de l'Homme
        # ======================================================
        print("\n📦 CNDH — Conseil National des Droits de l'Homme")
        marche_cndh = Marche(nom="CNDH", client="Conseil National des Droits de l'Homme")
        db.add(marche_cndh)
        db.flush()

        import xlrd
        cndh_path = os.path.join(MASTER_DATA_PATH, 'CNDH')
        cndh_files = {
            # Groupe 1 : (fichier, ville, checklist_type)
            "ERRACHIDIA S2 OK.XLS": ("Errachidia", ChecklistTypeEnum.CNDH_G1),
            "FES S2 OK.XLS": ("Fès", ChecklistTypeEnum.CNDH_G1),
            "GUELMIM S2 OK.XLS": ("Guelmim", ChecklistTypeEnum.CNDH_G1),
            "LAAYOUNE OK S2.XLS": ("Laâyoune", ChecklistTypeEnum.CNDH_G1),
            "MARRAKECH S2 OK.XLS": ("Marrakech", ChecklistTypeEnum.CNDH_G1),
            "OUJDA S2 MODIFIER.XLS": ("Oujda", ChecklistTypeEnum.CNDH_G1),
            "TANGER TETOUAN AL HOCEIMA S2 MODIFIER OK.XLS": ("Tanger-Tétouan-AH", ChecklistTypeEnum.CNDH_G1),
            "DAKHLA S2 OK.XLS": ("Dakhla", ChecklistTypeEnum.CNDH_G1),
            # Groupe 2
            "AGADIR ok S2.XLS": ("Agadir", ChecklistTypeEnum.CNDH_G2),
            "BENI MELLAL S2 OK.XLS": ("Béni Mellal", ChecklistTypeEnum.CNDH_G2),
            "CASABLANCA S2 OK.XLS": ("Casablanca", ChecklistTypeEnum.CNDH_G2),
        }

        for fname_cndh, (ville_cndh, ctype_cndh) in cndh_files.items():
            fpath = os.path.join(cndh_path, fname_cndh)
            if not os.path.exists(fpath):
                print(f"  ⚠️  {fname_cndh} non trouvé")
                continue
            sites_cfg = {"__default__": {"nom": f"CNDH {ville_cndh}", "ville": ville_cndh}}
            for sh_try in ['A', 'Feuil1']:
                sites_cfg[sh_try] = sites_cfg["__default__"]
            seed_excel_simple(db, marche_cndh.id, fpath, sites_cfg, ctype_cndh)

        # Siège Rabat (multi-feuilles)
        siege_path = None
        for f in os.listdir(cndh_path):
            if 'i' in f.lower() and 'ge' in f.lower() and f.endswith('.XLS'):
                siege_path = os.path.join(cndh_path, f)
                break

        if siege_path and os.path.exists(siege_path):
            site_siege = Site(nom="CNDH Siège Rabat", ville="Rabat", marche_id=marche_cndh.id,
                              checklist_type=ChecklistTypeEnum.CNDH_SIEGE,
                              feuilles=["SIEGE", "IFHD", "AGDAL"])
            db.add(site_siege)
            db.flush()

            wb_siege = xlrd.open_workbook(siege_path)
            for sh_name_s in ['SIEGE', 'IFHD', 'AGDAL']:
                if sh_name_s not in wb_siege.sheet_names():
                    continue
                ws_s = wb_siege.sheet_by_name(sh_name_s)
                rows_s = [ws_s.row_values(i) for i in range(ws_s.nrows)]
                h_idx_s = None
                headers_s = []
                for i, row in enumerate(rows_s):
                    row_str = ' '.join([str(c).lower() for c in row if c])
                    if 'serie' in row_str or 'article' in row_str:
                        headers_s = [str(c).strip() for c in row]
                        h_idx_s = i
                        break
                if h_idx_s is None:
                    continue
                count_s = 0
                for row in rows_s[h_idx_s + 1:]:
                    if not any(c for c in row):
                        continue
                    data_s = dict(zip(headers_s, row))
                    nserie_s = None
                    for key in headers_s:
                        if 'serie' in key.lower() or 'série' in key.lower():
                            nserie_s = clean(data_s.get(key))
                            break
                    if not nserie_s:
                        continue
                    article_s = None
                    for key in headers_s:
                        if 'article' in key.lower():
                            article_s = clean(data_s.get(key))
                            break
                    marque_s = None
                    for key in headers_s:
                        if 'marque' in key.lower():
                            marque_s = clean(data_s.get(key))
                            break
                    modele_s = None
                    for key in headers_s:
                        if 'modèle' in key.lower() or 'modele' in key.lower():
                            modele_s = clean(data_s.get(key))
                            break
                    empl_s = None
                    for key in headers_s:
                        if 'emplace' in key.lower():
                            empl_s = clean(data_s.get(key))
                            break

                    equip = Equipement(
                        nom=f"{marque_s or ''} {modele_s or ''}".strip() or article_s or 'Équipement',
                        designation=article_s,
                        type_equipement=get_type(article_s or ''),
                        site_id=site_siege.id,
                        marque=marque_s,
                        modele=modele_s,
                        numero_serie=nserie_s,
                        emplacement=empl_s,
                        sous_site=sh_name_s,
                    )
                    db.add(equip)
                    count_s += 1
                print(f"  ✅ CNDH Siège — {sh_name_s} : {count_s} équipements")

        # ======================================================
        # COMMIT FINAL
        # ======================================================
        db.commit()
        print("\n🎉 Seeding terminé avec succès !")
        print(f"   Marchés créés : {db.query(Marche).count()}")
        print(f"   Sites créés   : {db.query(Site).count()}")
        print(f"   Équipements   : {db.query(Equipement).count()}")

    except Exception as e:
        db.rollback()
        print(f"\n❌ Erreur : {e}")
        import traceback
        traceback.print_exc()
    finally:
        db.close()


if __name__ == "__main__":
    main()
