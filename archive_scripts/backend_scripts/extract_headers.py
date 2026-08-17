"""
Script d'extraction des numéros de marché et infos d'en-tête depuis les fichiers MASTER DATA
et mise à jour en base de données MySQL.
"""
import sys
import os
import re

# Fix Windows console encoding
sys.stdout.reconfigure(encoding='utf-8')


from app.db.database import engine
from sqlalchemy.orm import Session
from sqlalchemy import text
import openpyxl
from docx import Document

# ======================================================================
# CONFIG : Association CLIENT_KEY -> Nom du marché en base de données
# ======================================================================
# Cherche un marché dont le nom contient le mot-clé
MARCHE_KEYWORDS = {
    "ADM":       {"keywords": ["ADM"], "numero": "M0103/25",  "description": "Agence du Domaine de l'État"},
    "AMEE":      {"keywords": ["AMEE"], "numero": None,       "description": "Agence Marocaine pour l'Efficacité Energétique"},
    "ANCFCC":    {"keywords": ["ANCFCC"], "numero": None,     "description": "Agence Nationale de Conservation Foncière, du Cadastre et de la Cartographie"},
    "ANP":       {"keywords": ["ANP"], "numero": None,        "description": "Agence Nationale des Ports"},
    "AOH":       {"keywords": ["AOH"], "numero": None,        "description": "Agence de l'Oriental"},
    "CNDH":      {"keywords": ["CNDH"], "numero": None,       "description": "Conseil National des Droits de l'Homme"},
    "INPPLC":    {"keywords": ["INPPLC"], "numero": None,     "description": "Instance Nationale de la Probité, de la Prévention et de la Lutte contre la Corruption"},
    "MARSA_MAROC": {"keywords": ["MARSA"], "numero": None,    "description": "Marsa Maroc"},
    "MHAI":      {"keywords": ["MHAI", "HABOUS"], "numero": None, "description": "Ministère des Habous et Affaires Islamiques"},
    "MSANTE":    {"keywords": ["MSANTE", "SANTE"], "numero": None, "description": "Ministère de la Santé"},
    "ONEE":      {"keywords": ["ONEE"], "numero": None,       "description": "Office National de l'Électricité et de l'Eau Potable"},
    "ONP":       {"keywords": ["ONP"], "numero": None,        "description": "Office National des Pêches"},
}

MASTER_DATA_DIR = r"c:\Users\hanae\Desktop\Stage PFA 2026\MASTER DATA"


def extract_header_from_excel(filepath):
    """Extrait le numéro de marché et les infos d'en-tête depuis un fichier Excel."""
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        sheet = wb.active

        header_lines = []
        numero = None

        # Lire les 10 premières lignes
        for i, row in enumerate(sheet.iter_rows(max_row=10, values_only=True)):
            for cell in row:
                if cell and str(cell).strip():
                    text = str(cell).strip()
                    header_lines.append(text)
                    # Chercher un numéro de marché
                    m = re.search(r"March[ée]\s*[Nn]°?\s*([\w\/\.\-]+)", text)
                    if m and not numero:
                        numero = m.group(1).strip()
            if i >= 9:
                break

        wb.close()
        infos = " | ".join(header_lines[:5]) if header_lines else None
        return numero, infos

    except Exception as e:
        print(f"  [ERREUR Excel] {filepath}: {e}")
        return None, None


def extract_header_from_docx(filepath):
    """Extrait le numéro de marché et les infos d'en-tête depuis un fichier Word."""
    try:
        doc = Document(filepath)
        numero = None
        header_lines = []

        for para in doc.paragraphs[:20]:
            text = para.text.strip()
            if text:
                header_lines.append(text)
                m = re.search(r"March[ée]\s*[Nn]°?\s*([\w\/\.\-]+)", text)
                if m and not numero:
                    numero = m.group(1).strip()

        # Chercher aussi dans les tableaux (première cellule)
        for table in doc.tables[:1]:
            for row in table.rows[:5]:
                for cell in row.cells[:3]:
                    text = cell.text.strip()
                    if text:
                        m = re.search(r"March[ée]\s*[Nn]°?\s*([\w\/\.\-]+)", text)
                        if m and not numero:
                            numero = m.group(1).strip()

        infos = " | ".join(header_lines[:4]) if header_lines else None
        return numero, infos

    except Exception as e:
        print(f"  [ERREUR Docx] {filepath}: {e}")
        return None, None


def find_marche_files(client_key):
    """Trouve les fichiers d'un client dans MASTER DATA."""
    client_dir = os.path.join(MASTER_DATA_DIR, client_key)
    if not os.path.exists(client_dir):
        # Essayer avec espaces
        for folder in os.listdir(MASTER_DATA_DIR):
            if client_key.replace("_", " ").upper() in folder.upper():
                client_dir = os.path.join(MASTER_DATA_DIR, folder)
                break
        else:
            return []

    files = []
    for f in os.listdir(client_dir):
        if f.startswith("~$"):
            continue
        fpath = os.path.join(client_dir, f)
        if f.lower().endswith((".xlsx", ".xls", ".xlsm")):
            files.append(("excel", fpath))
        elif f.lower().endswith(".docx"):
            files.append(("docx", fpath))
    return files


def main():
    with Session(engine) as session:
        # Lire tous les marchés existants
        result = session.execute(text("SELECT id, nom FROM marches")).fetchall()
        marches = [(row[0], row[1]) for row in result]
        print(f"✅ {len(marches)} marché(s) trouvé(s) en base de données :")
        for mid, mnom in marches:
            print(f"   - [{mid}] {mnom}")

        print("\n========== Extraction des en-têtes depuis MASTER DATA ==========\n")

        updates = []

        for client_key, cfg in MARCHE_KEYWORDS.items():
            # Trouver le marché correspondant
            marche_id = None
            for mid, mnom in marches:
                for kw in cfg["keywords"]:
                    if kw.upper() in mnom.upper():
                        marche_id = mid
                        break
                if marche_id:
                    break

            if not marche_id:
                print(f"⚠️  [{client_key}] Aucun marché trouvé en base pour les keywords {cfg['keywords']}")
                continue

            print(f"\n📂 [{client_key}] Marché ID={marche_id}")

            # Trouver le premier fichier
            files = find_marche_files(client_key)
            if not files:
                print(f"   ⚠️  Aucun fichier trouvé dans MASTER DATA/{client_key}")
                # Utiliser les valeurs pré-définies si disponibles
                if cfg.get("numero"):
                    updates.append((marche_id, cfg["numero"], cfg.get("description"), client_key))
                continue

            # Prendre le premier fichier Word ou Excel
            filetype, filepath = files[0]
            print(f"   📄 Fichier : {os.path.basename(filepath)}")

            numero = None
            infos = None

            if filetype == "excel":
                numero, infos = extract_header_from_excel(filepath)
            elif filetype == "docx":
                numero, infos = extract_header_from_docx(filepath)

            # Fallback : utiliser les valeurs pré-définies
            if not numero and cfg.get("numero"):
                numero = cfg["numero"]
                print(f"   ℹ️  Utilisation du numéro pré-défini: {numero}")
            elif numero:
                print(f"   ✅ Numéro extrait: {numero}")

            if not infos:
                infos = cfg.get("description")
            else:
                print(f"   ✅ Infos extraites: {infos[:80]}...")

            updates.append((marche_id, numero, infos, client_key))

        print("\n\n========== Mise à jour de la base de données ==========\n")
        for marche_id, numero, infos, client_key in updates:
            session.execute(
                text("UPDATE marches SET numero = :numero, informations_entete = :infos WHERE id = :id"),
                {"numero": numero, "infos": infos, "id": marche_id}
            )
            print(f"✅ [{client_key}] ID={marche_id} → N°={numero} | Infos={str(infos)[:60] if infos else 'N/A'}...")

        session.commit()
        print(f"\n🎉 {len(updates)} marché(s) mis à jour avec succès !")


if __name__ == "__main__":
    main()
