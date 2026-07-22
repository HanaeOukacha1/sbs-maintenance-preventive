"""
Script d'analyse des fichiers Word (.docx) du MASTER DATA SBS.
Extrait les tableaux et textes pour comprendre la structure des checklists.
"""
import os
import json

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    import subprocess, sys
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

MASTER_DATA_PATH = r"c:\Users\hanae\Desktop\Stage PFA 2026\MASTER DATA"
OUTPUT_PATH = r"c:\Users\hanae\Desktop\Stage PFA 2026\MASTER DATA\analyse_words.json"

def extraire_tableaux(filepath: str) -> list:
    """Extrait tous les tableaux d'un fichier Word"""
    doc = Document(filepath)
    tableaux = []

    for i, table in enumerate(doc.tables):
        lignes = []
        for row in table.rows:
            cellules = [cell.text.strip() for cell in row.cells]
            # Dédupliquer les cellules fusionnées
            cellules_clean = []
            prev = None
            for c in cellules:
                if c != prev:
                    cellules_clean.append(c)
                prev = c
            if any(cellules_clean):
                lignes.append(cellules_clean)
        if lignes:
            tableaux.append({"tableau_index": i, "lignes": lignes})

    return tableaux

def extraire_texte(filepath: str) -> list:
    """Extrait les paragraphes non vides d'un fichier Word"""
    doc = Document(filepath)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# Analyse de tous les fichiers Word
rapport = {}

print("=" * 60)
print("ANALYSE DES FICHIERS WORD SBS")
print("=" * 60)

for marche_nom in sorted(os.listdir(MASTER_DATA_PATH)):
    marche_path = os.path.join(MASTER_DATA_PATH, marche_nom)
    if not os.path.isdir(marche_path):
        continue

    fichiers_word = [f for f in os.listdir(marche_path) if f.lower().endswith(".docx")]
    if not fichiers_word:
        continue

    print(f"\n📁 MARCHÉ : {marche_nom}")
    rapport[marche_nom] = []

    for fichier in sorted(fichiers_word):
        filepath = os.path.join(marche_path, fichier)
        print(f"  📝 Lecture : {fichier}")

        try:
            tableaux = extraire_tableaux(filepath)
            textes = extraire_texte(filepath)

            info = {
                "fichier": fichier,
                "nb_tableaux": len(tableaux),
                "textes_cles": textes[:10],  # 10 premiers paragraphes
                "tableaux": tableaux[:3]      # 3 premiers tableaux max
            }
            rapport[marche_nom].append(info)

            print(f"    → {len(tableaux)} tableau(x) trouvé(s)")
            if tableaux:
                print(f"    → 1er tableau ({len(tableaux[0]['lignes'])} lignes)")
                for l in tableaux[0]["lignes"][:5]:
                    print(f"       {l}")

        except Exception as e:
            print(f"    ⚠️  Erreur : {e}")
            rapport[marche_nom].append({"fichier": fichier, "erreur": str(e)})

with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
    json.dump(rapport, f, ensure_ascii=False, indent=2)

print(f"\n✅ Rapport sauvegardé : {OUTPUT_PATH}")
